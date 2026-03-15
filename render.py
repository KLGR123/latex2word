#!/usr/bin/env python3
"""
render.py — Render translated JSON paragraphs into a Word (.docx) document.

Reads the replaced.json pipeline output and appends each chunk to an existing
or newly created .docx file using python-docx.

Design decisions:
  - All LaTeX text (body, inline math, display equations) is converted to
    native Word OMML math via pandoc (-f latex -t docx); XML nodes are
    injected directly — no lossy plain-text math rendering.
  - Footnotes: pandoc's word/footnotes.xml is merged into the document with
    remapped IDs so footnote content is preserved.
  - Side-by-side figures (minipage / subfigure) are laid out in a borderless
    Word table so images and their captions stay aligned.
  - Tables are kept as raw LaTeX source for now (pending a dedicated renderer).
  - Fonts use all four XML slots (ascii / eastAsia / hAnsi / cs):
      body    -> 宋体      heading -> 黑体
      caption -> 楷体      code    -> Courier New

Usage:
    python render.py --json replaced.json --docx output.docx
                     [--figures-dir .] [--citations citations.json]
                     [--skip-title]

Dependencies (Python): python-docx, lxml, cn2an
Dependencies (system): pandoc
"""

import argparse
import copy
import json
import os
import re
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree

from table import render_latex_table_to_docx

# ─────────────────────────────────────────────────────────────────────────────
# Optional dependency: cn2an (Arabic -> Chinese numerals for chapter titles)
# ─────────────────────────────────────────────────────────────────────────────

try:
    import cn2an as _cn2an
    def _to_chinese_number(n: int) -> str:
        return _cn2an.an2cn(str(n))
except ImportError:
    def _to_chinese_number(n: int) -> str:  # type: ignore[misc]
        return str(n)


# ─────────────────────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────────────────────

FONT_BODY    = "宋体"
FONT_HEADING = "黑体"
FONT_CAPTION = "楷体"
FONT_CODE    = "Courier New"

SIZE_BODY    = 12
SIZE_H1      = 16
SIZE_H2      = 14
SIZE_H3      = 12
SIZE_CAPTION = 10
SIZE_CODE    = 9

BG_THEOREM = "EBF2FA"   # light blue  -- theorem / lemma / definition blocks
BG_CODE    = "F5F5F5"   # light grey  -- verbatim code blocks
BG_TABLE   = "FFF3CD"   # light amber -- raw-LaTeX table placeholder

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"

MATH_ENV_LABELS = {
    "definition":  "定义",
    "lemma":       "引理",
    "theorem":     "定理",
    "corollary":   "推论",
    "proposition": "命题",
    "remark":      "注记",
    "claim":       "断言",
    "example":     "例",
    "fact":        "事实",
}

_TEX_PREAMBLE  = (
    "\\documentclass{article}\n"
    "\\usepackage{amsmath,amssymb,amsfonts,mathtools,bm}\n"
    "\\begin{document}\n"
)
_TEX_POSTAMBLE = "\n\\end{document}\n"

# Registry of temporary files created during a run (e.g. PDF→PNG conversions).
# Populated by _pdf_to_image(); cleaned up at the end of main().
_g_temp_files: List[str] = []


# ─────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def _make_rFonts(font_name: str) -> Any:
    """Return a <w:rFonts> element with all four font slots set to font_name."""
    el = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:eastAsia", "w:hAnsi", "w:cs"):
        el.set(qn(attr), font_name)
    return el


def set_run_font(run, font_name: str, size_pt: float = None,
                 bold: bool = False, italic: bool = False,
                 color: RGBColor = None) -> None:
    """Apply font_name to a python-docx Run across all four w:rFonts slots."""
    rPr = run._r.get_or_add_rPr()
    old = rPr.find(qn("w:rFonts"))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, _make_rFonts(font_name))
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color


def _set_para_font(para, font_name: str, size_pt: float) -> None:
    """Set paragraph-level default font so runs without explicit rPr inherit it."""
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    old = rPr.find(qn("w:rFonts"))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, _make_rFonts(font_name))
    half_pts = str(int(size_pt * 2))
    for tag in ("w:sz", "w:szCs"):
        e = rPr.find(qn(tag))
        if e is None:
            e = OxmlElement(tag)
            rPr.append(e)
        e.set(qn("w:val"), half_pts)


def _fix_lxml_run_fonts(r_elem, font_name: str, size_pt: float) -> None:
    """Overwrite w:rFonts on a raw lxml <w:r> element (from pandoc output)."""
    rPr = r_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r_elem.insert(0, rPr)
    old = rPr.find(qn("w:rFonts"))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, _make_rFonts(font_name))
    half_pts = str(int(size_pt * 2))
    for tag in ("w:sz", "w:szCs"):
        e = rPr.find(qn(tag))
        if e is None:
            e = OxmlElement(tag)
            rPr.append(e)
        e.set(qn("w:val"), half_pts)


def _fix_math_text_runs(elem, font_name: str) -> None:
    """
    Set font on <m:r> runs that carry <m:nor> (i.e. \\text{} inside math),
    so that text-mode content inside formulas uses the document body font.
    """
    m_r   = f"{{{NS_M}}}r"
    m_rPr = f"{{{NS_M}}}rPr"
    m_nor = f"{{{NS_M}}}nor"
    for run in elem.iter(m_r):
        rPr = run.find(m_rPr)
        if rPr is None or rPr.find(m_nor) is None:
            continue
        w_rPr = run.find(qn("w:rPr"))
        if w_rPr is None:
            w_rPr = OxmlElement("w:rPr")
            run.insert(0, w_rPr)
        old = w_rPr.find(qn("w:rFonts"))
        if old is not None:
            w_rPr.remove(old)
        w_rPr.insert(0, _make_rFonts(font_name))


# ─────────────────────────────────────────────────────────────────────────────
# Paragraph formatting helpers
# ─────────────────────────────────────────────────────────────────────────────

def _spacing(para, before_pt: float = 0.0, after_pt: float = 6.0) -> None:
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(int(before_pt * 20)))
    sp.set(qn("w:after"),  str(int(after_pt  * 20)))
    para._p.get_or_add_pPr().append(sp)


def _indent(para, left_cm: float = 0.0, first_line_cm: float = 0.0,
            hanging_cm: float = 0.0) -> None:
    ind = OxmlElement("w:ind")
    if left_cm:
        ind.set(qn("w:left"),      str(int(left_cm       * 567)))
    if first_line_cm:
        ind.set(qn("w:firstLine"), str(int(first_line_cm * 567)))
    if hanging_cm:
        ind.set(qn("w:hanging"),   str(int(hanging_cm    * 567)))
    para._p.get_or_add_pPr().append(ind)


def _shade(para, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    para._p.get_or_add_pPr().append(shd)


# ─────────────────────────────────────────────────────────────────────────────
# Space-cleaning helper
# ─────────────────────────────────────────────────────────────────────────────

def _clean_spaces(text: str) -> str:
    """
    Remove spaces that are not between two ASCII printable non-space chars.
    Keeps spaces inside English words / between digits; strips spaces adjacent
    to CJK characters that the LLM translator inserts unnecessarily.
    """
    if not text:
        return text
    chars = list(text)
    result: List[str] = []
    for i, ch in enumerate(chars):
        if ch != " ":
            result.append(ch)
            continue
        prev_ch = chars[i - 1] if i > 0 else "\0"
        next_ch = chars[i + 1] if i < len(chars) - 1 else "\0"
        if (prev_ch.isascii() and prev_ch.isprintable() and not prev_ch.isspace()
                and next_ch.isascii() and next_ch.isprintable() and not next_ch.isspace()):
            result.append(ch)
    return "".join(result)


# ─────────────────────────────────────────────────────────────────────────────
# Footnote management
# ─────────────────────────────────────────────────────────────────────────────

_g_footnote_next_id: int = 1
_g_footnotes_root: Optional[Any] = None   # lxml element for the footnotes part

_FN_TAG        = f"{{{NS_W}}}footnote"
_FN_REF_TAG    = f"{{{NS_W}}}footnoteReference"
_FN_ID_ATTR    = f"{{{NS_W}}}id"
_FN_TYPE_ATTR  = f"{{{NS_W}}}type"
_FN_SKIP_TYPES = frozenset({"separator", "continuationSeparator"})

# Minimal footnotes.xml skeleton required by the OOXML spec.
# IDs -1 and 0 are the mandatory separator / continuationSeparator entries.
_FN_SKELETON_XML = (
    '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:footnote w:id="-1" w:type="separator">'
    '<w:p><w:r><w:separator/></w:r></w:p>'
    '</w:footnote>'
    '<w:footnote w:id="0" w:type="continuationSeparator">'
    '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
    '</w:footnote>'
    '</w:footnotes>'
)


def _init_footnotes(doc: Document) -> None:
    """
    Ensure the document has a footnotes part, creating one from scratch when
    absent (freshly created .docx files have no footnotes relationship).
    Records the next available footnote ID so that footnotes from pandoc
    fragments are assigned non-conflicting IDs.
    """
    global _g_footnote_next_id, _g_footnotes_root

    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    from docx.opc.part import XmlPart
    from docx.oxml import parse_xml

    part = doc.part

    # Find an existing footnotes relationship first.
    fn_element = None
    for rel in part.rels.values():
        if rel.reltype == RT.FOOTNOTES:
            fn_element = rel.target_part._element
            break

    if fn_element is None:
        # Create a brand-new footnotes part and register the relationship.
        fn_partname    = PackURI("/word/footnotes.xml")
        fn_contenttype = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.footnotes+xml"
        )
        fn_element  = parse_xml(_FN_SKELETON_XML)
        fn_opc_part = XmlPart(fn_partname, fn_contenttype, fn_element, part.package)
        part.relate_to(fn_opc_part, RT.FOOTNOTES)
        print("  [footnotes] Created new footnotes part.")

    _g_footnotes_root = fn_element
    max_id = max(
        (int(fn.get(_FN_ID_ATTR, "0"))
         for fn in _g_footnotes_root.findall(_FN_TAG)
         if fn.get(_FN_TYPE_ATTR, "") not in _FN_SKIP_TYPES),
        default=0,
    )
    _g_footnote_next_id = max(max_id + 1, 1)
    print(f"  [footnotes] Initialised -- next ID = {_g_footnote_next_id}.")


def _fix_footnote_fonts(fn_elem) -> None:
    """
    Set FONT_BODY (宋体) on every <w:r> run inside a footnote element.
    Pandoc does not know the document font, so footnote runs inherit whatever
    default the skeleton XML provides; this call normalises them all.
    """
    fn_size = SIZE_BODY - 1          # 11 pt — standard footnote size
    for r_elem in fn_elem.iter(f"{{{NS_W}}}r"):
        _fix_lxml_run_fonts(r_elem, FONT_BODY, fn_size)


def _merge_pandoc_footnotes(doc_tree, fn_xml_bytes: Optional[bytes]) -> None:
    """
    Remap footnote IDs from a pandoc fragment and merge them into the document.
    Modifies doc_tree in-place (updates <w:footnoteReference @w:id> values).
    """
    global _g_footnote_next_id, _g_footnotes_root
    if fn_xml_bytes is None or _g_footnotes_root is None:
        return
    try:
        fn_tree = etree.fromstring(fn_xml_bytes)
    except Exception as exc:
        print(f"  [footnotes] Could not parse footnotes.xml: {exc}")
        return

    pandoc_fns: Dict[int, Any] = {
        int(fn.get(_FN_ID_ATTR, "")): fn
        for fn in fn_tree.findall(_FN_TAG)
        if fn.get(_FN_TYPE_ATTR, "") not in _FN_SKIP_TYPES
        and fn.get(_FN_ID_ATTR, "").lstrip("-").isdigit()
    }
    if not pandoc_fns:
        return

    id_map = {old: _g_footnote_next_id + i for i, old in enumerate(sorted(pandoc_fns))}
    _g_footnote_next_id += len(id_map)

    for ref in doc_tree.iter(_FN_REF_TAG):
        try:
            new_id = id_map.get(int(ref.get(_FN_ID_ATTR, "")))
            if new_id is not None:
                ref.set(_FN_ID_ATTR, str(new_id))
        except (ValueError, TypeError):
            pass

    for old_id, new_id in id_map.items():
        fn_elem = copy.deepcopy(pandoc_fns[old_id])
        fn_elem.set(_FN_ID_ATTR, str(new_id))
        _fix_footnote_fonts(fn_elem)          # ensure Chinese text uses 宋体
        _g_footnotes_root.append(fn_elem)


# ─────────────────────────────────────────────────────────────────────────────
# pandoc bridge: LaTeX -> Word XML
# ─────────────────────────────────────────────────────────────────────────────

def _pandoc_to_xml(latex_body: str) -> Tuple[Optional[bytes], Optional[bytes]]:
    """
    Compile latex_body through pandoc and return (document.xml, footnotes.xml).
    Either value is None on failure or absence.
    """
    src = _TEX_PREAMBLE + latex_body + _TEX_POSTAMBLE
    tex_path = docx_path = None
    try:
        with tempfile.NamedTemporaryFile(
            suffix=".tex", mode="w", encoding="utf-8", delete=False
        ) as f:
            f.write(src)
            tex_path = f.name
        docx_path = tex_path.replace(".tex", ".docx")

        result = subprocess.run(
            ["pandoc", tex_path, "-f", "latex", "-t", "docx",
             "-o", docx_path, "--quiet"],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            print(f"  [pandoc] {result.stderr[:200]}")
            return None, None

        with zipfile.ZipFile(docx_path) as z:
            names = z.namelist()
            return (
                z.read("word/document.xml"),
                z.read("word/footnotes.xml") if "word/footnotes.xml" in names else None,
            )
    except Exception as exc:
        print(f"  [pandoc] {exc}")
        return None, None
    finally:
        for p in (tex_path, docx_path):
            if p:
                try:
                    os.unlink(p)
                except OSError:
                    pass


def _inject_pandoc_paras(latex_body: str, doc,
                          font_name: str, size_pt: float,
                          align: WD_ALIGN_PARAGRAPH,
                          indent_first_cm: float = 0.0,
                          indent_left_cm: float = 0.0,
                          before_pt: float = 0.0,
                          after_pt: float = 6.0) -> bool:
    """
    Convert latex_body via pandoc and inject each non-empty <w:p> into doc.
    Overwrites all run fonts to font_name/size_pt and merges footnotes.
    Returns True if at least one paragraph was injected.
    """
    doc_xml, fn_xml = _pandoc_to_xml(latex_body)
    if doc_xml is None:
        return False

    tree = etree.fromstring(doc_xml)
    # Merge footnotes before copying nodes so remapped IDs are already in tree.
    _merge_pandoc_footnotes(tree, fn_xml)

    injected = 0
    for src_p in tree.findall(f".//{{{NS_W}}}p"):
        # Skip structurally empty paragraphs (only a <w:pPr> child)
        if all(etree.QName(c.tag).localname == "pPr" for c in src_p):
            continue

        dst = doc.add_paragraph()
        dst.alignment = align
        _indent(dst, left_cm=indent_left_cm, first_line_cm=indent_first_cm)
        _spacing(dst, before_pt=before_pt, after_pt=after_pt)
        _set_para_font(dst, font_name, size_pt)

        for child in src_p:
            if etree.QName(child.tag).localname == "pPr":
                continue
            node = copy.deepcopy(child)
            if etree.QName(node.tag).localname == "r":
                _fix_lxml_run_fonts(node, font_name, size_pt)
            _fix_math_text_runs(node, font_name)
            dst._p.append(node)

        injected += 1

    return injected > 0


# ─────────────────────────────────────────────────────────────────────────────
# LaTeX utility helpers
# ─────────────────────────────────────────────────────────────────────────────

def _clean_cross_refs(text: str) -> str:
    """Remove doubled section label artefacts, e.g. '第3.5节 节' -> '第3.5节'."""
    return re.sub(r'(第[0-9.]+[节章])\s*节', r'\1', text)


def _extract_caption(text: str) -> str:
    """Return the content of the first \\caption{...} in text, or empty string."""
    m = re.search(r'\\caption\{((?:[^{}]|\{[^{}]*\})*)\}', text, re.DOTALL)
    return m.group(1).strip() if m else ""


def _extract_env_body(text: str, env_name: str) -> Tuple[str, str]:
    """
    Extract (optional_arg, body) from \\begin{env}[opt]...\\end{env}.
    Returns ("", text) when the environment is not found.
    """
    m = re.search(
        r'\\begin\{' + re.escape(env_name) + r'\*?\}'
        r'(?:\[([^\]]*)\])?'
        r'(.*?)'
        r'\\end\{' + re.escape(env_name) + r'\*?\}',
        text, re.DOTALL | re.IGNORECASE,
    )
    return ((m.group(1) or "").strip(), m.group(2).strip()) if m else ("", text)


# ─────────────────────────────────────────────────────────────────────────────
# Core paragraph renderer
# ─────────────────────────────────────────────────────────────────────────────

def _render_para(doc, latex_text: str,
                 font_name: str = FONT_BODY,
                 size_pt: float = SIZE_BODY,
                 align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
                 indent_first_cm: float = 0.74,
                 indent_left_cm: float = 0.0,
                 before_pt: float = 0.0,
                 after_pt: float = 6.0) -> bool:
    """
    Render LaTeX/Chinese text into doc via pandoc (inline math -> OMML).
    Falls back to a stripped plain-text paragraph on pandoc failure.
    Returns True if pandoc succeeded.
    """
    text = _clean_cross_refs(latex_text).strip()
    if not text:
        return True

    ok = _inject_pandoc_paras(
        text, doc,
        font_name=font_name, size_pt=size_pt, align=align,
        indent_first_cm=indent_first_cm, indent_left_cm=indent_left_cm,
        before_pt=before_pt, after_pt=after_pt,
    )
    if ok:
        return True

    # Strip LaTeX markup as a best-effort plain-text fallback
    plain = re.sub(r'\\[a-zA-Z]+\*?\{([^{}]*)\}', r'\1', text)
    plain = re.sub(r'\\[a-zA-Z]+\*?', '', plain)
    plain = plain.replace('$', '').replace('{', '').replace('}', '').strip()
    p = doc.add_paragraph()
    p.alignment = align
    _indent(p, left_cm=indent_left_cm, first_line_cm=indent_first_cm)
    _spacing(p, before_pt=before_pt, after_pt=after_pt)
    set_run_font(p.add_run(plain), font_name, size_pt)
    return False


# ─────────────────────────────────────────────────────────────────────────────
# Chunk classifier — environment regexes
# Each pattern matches \begin{env_name} at the start of a stripped chunk.
# ─────────────────────────────────────────────────────────────────────────────

_RE_SECTION       = re.compile(r'\\section\*?\{')
_RE_SUBSECTION    = re.compile(r'\\subsection\*?\{')
_RE_SUBSUBSECTION = re.compile(r'\\subsubsection\*?\{')
_RE_ABSTRACT      = re.compile(r'\\begin\{abstract\}', re.IGNORECASE)

_RE_FIGURE = re.compile(
    r'\\begin\{'
    r'(?:figure\*?'
    r'|subfigure\*?'
    r'|subfloat'
    r'|wrapfigure\*?'
    r'|SCfigure\*?'
    r'|floatrow'
    r'|ffigbox'
    r'|capbeside'
    r'|sidewaysfigure\*?'
    r'|turnfigure'
    r'|captionedbox'
    r'|tikzpicture'
    r'|tikzfigure'
    r'|floatingfigure'
    r'|figwindow'
    r'|cutout'
    r'|overpic'
    r'|teaserfigure'
    r'|figurehere'
    r')\}'
)

_RE_TABLE = re.compile(
    r'\\begin\{'
    r'(?:table\*?'
    r'|tabular\*?'
    r'|tabularx\*?'
    r'|tabulary'
    r'|tabularray'
    r'|tblr'
    r'|longtblr'
    r'|longtable\*?'
    r'|supertabular\*?'
    r'|mpsupertabular'
    r'|xtabular'
    r'|sidewaystable\*?'
    r'|sidewaystabular'
    r'|turntable'
    r'|threeparttable'
    r'|threeparttablex'
    r'|tabu'
    r'|longtabu'
    r'|tabbing'
    r'|wraptable\*?'
    r'|ttabbox'
    r'|ctable'
    r'|array'
    r'|spreadtab'
    r'|tablehere'
    r'|adjustbox'
    r')\}'
)

_RE_ALGORITHM = re.compile(
    r'\\begin\{'
    r'(?:algorithm\*?'
    r'|algorithmic'
    r'|algorithmicx'
    r'|algpseudocode'
    r'|algorithm2e\*?'
    r'|pseudocode'
    r'|lstpseudocode'
    r'|myalgorithm'
    r'|algo'
    r'|proc'
    r'|procedure'
    r'|function'
    r')\}'
)

_RE_VERBATIM = re.compile(
    r'\\begin\{'
    r'(?:verbatim\*?'
    r'|lstlisting'
    r'|minted\*?'
    r'|Verbatim\*?'
    r'|BVerbatim'
    r'|LVerbatim'
    r'|SaveVerbatim'
    r'|alltt'
    r'|verbatimtab'
    r'|listing'
    r'|tcolorbox'
    r'|tcblisting'
    r'|mdframed'
    r'|spverbatim'
    r'|codeblock'
    r'|codebox'
    r'|sourcecode'
    r'|pycode'
    r'|bashcode'
    r'|jsoncode'
    r'|xmlcode'
    r'|sqlcode'
    r'|exampleblock'
    r'|example'
    r')\}'
)

_RE_EQUATION = re.compile(
    r'\\begin\{'
    r'(?:equation\*?'
    r'|displaymath'
    r'|align\*?'
    r'|aligned'
    r'|alignat\*?'
    r'|alignedat'
    r'|gather\*?'
    r'|gathered'
    r'|multline\*?'
    r'|flalign\*?'
    r'|split'
    r'|cases\*?'
    r'|dcases\*?'
    r'|rcases\*?'
    r'|numcases'
    r'|subequations'
    r'|eqnarray\*?'
    r'|empheq'
    r'|dmath\*?'
    r'|dseries\*?'
    r'|dgroup\*?'
    r'|darray\*?'
    r'|IEEEeqnarray\*?'
    r'|math'
    r')\}'
)

_RE_MATHENV = re.compile(
    r'\\begin\{'
    r'(definition|lemma|theorem|corollary|proposition'
    r'|remark|claim|example|fact'
    r')\*?\}',
    re.IGNORECASE
)

_RE_PROOF = re.compile(r'\\begin\{proof\}', re.IGNORECASE)


def classify(para: dict) -> str:
    label = para.get("env_label", "")
    text  = para.get("text", "").strip()

    if label == "摘要":               return "abstract"
    if label.startswith("式"):        return "equation"
    if label.startswith("图"):        return "figure"
    if label.startswith("表"):        return "table"
    if label.startswith("代码"):      return "code"
    if label.startswith("算法"):      return "code"   # rendered same as code

    if _RE_ABSTRACT.match(text):      return "abstract"
    if _RE_SECTION.match(text):       return "section"
    if _RE_SUBSECTION.match(text):    return "subsection"
    if _RE_SUBSUBSECTION.match(text): return "subsubsection"
    if _RE_EQUATION.match(text):      return "equation"
    if _RE_FIGURE.match(text):        return "figure"
    if _RE_TABLE.match(text):         return "table"
    if _RE_ALGORITHM.match(text):     return "code"
    if _RE_VERBATIM.match(text):      return "code"
    if _RE_MATHENV.match(text):       return "mathenv"
    if _RE_PROOF.match(text):         return "proof"
    return "paragraph"


# ─────────────────────────────────────────────────────────────────────────────
# Renderers
# ─────────────────────────────────────────────────────────────────────────────

def render_chapter_title(doc, meta: dict) -> None:
    """Render the chapter heading; Arabic chapter number is converted to Chinese."""
    chapter  = meta.get("chapter", "")
    title_zh = meta.get("title_translation", meta.get("title", ""))
    title_zh = re.sub(r'^#+\s+', '', title_zh)   # strip leading # marks
    try:
        chapter_cn = _to_chinese_number(int(chapter))
    except (ValueError, TypeError):
        chapter_cn = str(chapter)
    h = doc.add_heading(f"第{chapter_cn}章  {title_zh}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        set_run_font(run, FONT_HEADING, SIZE_H1 + 2, bold=True)
    _spacing(h, before_pt=24, after_pt=12)


def _extract_heading_text(translation: str, cmd: str) -> Tuple[str, str]:
    """
    Parse \\cmd{heading text} and return (heading, trailing_body).
    Handles arbitrary brace nesting.
    """
    m = re.match(r'\\' + re.escape(cmd) + r'\*?\{', translation.strip())
    if not m:
        return "", translation
    depth, i = 0, m.end() - 1
    while i < len(translation):
        if   translation[i] == '{': depth += 1
        elif translation[i] == '}':
            depth -= 1
            if depth == 0:
                return translation[m.end():i].strip(), translation[i + 1:].strip()
        i += 1
    return translation[m.end():].strip(), ""


def render_heading(doc, para: dict, level: int) -> None:
    translation = para.get("translation", para.get("text", "")).strip()
    env_label   = para.get("env_label", "")
    cmd_map     = {1: "section", 2: "subsection", 3: "subsubsection"}

    heading_text, body_text = _extract_heading_text(translation, cmd_map[level])
    if not heading_text:
        m = re.search(r'\{([^}]+)\}', translation)
        heading_text = m.group(1) if m else translation
        body_text    = ""

    # Strip remaining LaTeX markup from the visible heading string
    heading_clean = re.sub(r'\\[a-zA-Z]+\*?\{([^{}]*)\}', r'\1', heading_text)
    heading_clean = re.sub(r'\\[a-zA-Z]+\*?', '', heading_clean).strip()
    heading_clean = _clean_spaces(heading_clean)
    heading_clean = re.sub(r'^#+\s+', '', heading_clean)   # strip leading # marks

    # Numeric prefix from env_label (e.g. "3.5" from "3.5节")
    numeric = re.sub(r'[节小].*$', '', env_label)
    prefix  = (numeric + "  ") if re.match(r'^\d', numeric) else ""

    h = doc.add_heading(prefix + heading_clean, level=level)
    for run in h.runs:
        set_run_font(run, FONT_HEADING, {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3}[level], bold=True)
    _spacing(h, before_pt=14, after_pt=5)

    if body_text.strip():
        _render_para(doc, _clean_spaces(body_text))


def render_abstract(doc, para: dict) -> None:
    translation = para.get("translation", para.get("text", ""))
    _, body     = _extract_env_body(translation, "abstract")
    if not body:
        body = translation
    h = doc.add_heading("摘要", level=1)
    for run in h.runs:
        set_run_font(run, FONT_HEADING, SIZE_H1, bold=True)
    _spacing(h, before_pt=14, after_pt=6)
    _render_para(doc, _clean_spaces(body), indent_first_cm=0.74, indent_left_cm=1.0)


def render_paragraph(doc, para: dict) -> None:
    _render_para(doc, _clean_spaces(para.get("translation", para.get("text", ""))))


def render_equation(doc, para: dict) -> None:
    translation = para.get("translation", para.get("text", "")).strip()
    env_label   = para.get("env_label", "")

    ok = _inject_pandoc_paras(
        _clean_cross_refs(translation), doc,
        font_name=FONT_BODY, size_pt=SIZE_BODY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before_pt=6.0, after_pt=2.0,
    )
    if not ok:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, before_pt=6, after_pt=2)
        run = p.add_run(translation)
        run.font.name = FONT_CODE
        run.font.size = Pt(SIZE_CODE)

    if env_label:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _spacing(lp, before_pt=0, after_pt=8)
        set_run_font(lp.add_run(f"（{env_label}）"), FONT_BODY, 10)


def _add_caption(doc, label: str, caption_latex: str) -> None:
    """Render a figure/table caption (楷体, centred), with inline-math support."""
    if not label and not caption_latex:
        return
    cleaned = _clean_spaces(caption_latex) if caption_latex else ""

    # Build full LaTeX string: bold label + space + caption text.
    # A plain space after the closing brace is the most portable separator —
    # \quad is sometimes silently dropped by pandoc when adjacent to CJK text.
    full_latex = (
        f"\\textbf{{{label}}} {cleaned}" if cleaned
        else f"\\textbf{{{label}}}"
    )

    # Primary path: pandoc handles both text and inline math together
    ok = _inject_pandoc_paras(
        full_latex, doc,
        font_name=FONT_CAPTION, size_pt=SIZE_CAPTION,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before_pt=2.0, after_pt=8.0,
    )
    if not ok:
        # Fallback: split on $...$ and handle math segments separately
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, before_pt=2, after_pt=8)
        parts = re.split(r'(\$[^$]+\$)', f"{label} {cleaned}")
        for part in parts:
            if part.startswith("$") and part.endswith("$"):
                # Attempt to render math inline via pandoc into a temp para,
                # then fall back to raw text if that also fails
                math_ok = _inject_pandoc_paras(
                    part, doc,
                    font_name=FONT_CAPTION, size_pt=SIZE_CAPTION,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                if not math_ok:
                    set_run_font(p.add_run(part), FONT_CAPTION, SIZE_CAPTION)
            else:
                set_run_font(p.add_run(part), FONT_CAPTION, SIZE_CAPTION, italic=False)


def _find_minipages(text: str) -> List[Dict[str, Any]]:
    """
    Extract per-subfigure data from minipage / subfigure / subfloat blocks.
    Returns a list of {"images": [...], "caption": "..."} dicts.
    """
    results = []
    pat = re.compile(
        r'\\begin\{(?:minipage|subfigure|subfloat)\}'
        r'(?:\[[^\]]*\])?(?:\{[^}]*\})?'
        r'(.*?)'
        r'\\end\{(?:minipage|subfigure|subfloat)\}',
        re.DOTALL,
    )
    for m in pat.finditer(text):
        content = m.group(1)
        imgs = re.findall(r'\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}', content)
        cap  = _extract_caption(content)
        if imgs or cap:
            results.append({"images": imgs, "caption": cap})
    return results


def _remove_table_borders(table) -> None:
    """Remove all visible borders from a python-docx table."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def _pdf_to_image(pdf_path: Path) -> Optional[Path]:
    """
    Rasterise the first page of a PDF to a temporary PNG file.

    Tries PyMuPDF (fitz) first, then falls back to pdf2image (poppler).
    The resulting path is registered in _g_temp_files for cleanup.
    Returns None if no suitable library is available or conversion fails.
    """
    global _g_temp_files

    # --- attempt 1: PyMuPDF (pip install pymupdf) ---
    try:
        import fitz  # type: ignore
        pdf_doc = fitz.open(str(pdf_path))
        page    = pdf_doc[0]
        mat     = fitz.Matrix(2.0, 2.0)   # 2× zoom → ~144 dpi effective
        pix     = page.get_pixmap(matrix=mat)
        tmp     = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        pix.save(tmp.name)
        tmp.close()
        pdf_doc.close()
        _g_temp_files.append(tmp.name)
        print(f"  [pdf→png] {pdf_path.name} -> {tmp.name} (fitz)")
        return Path(tmp.name)
    except ImportError:
        pass
    except Exception as exc:
        print(f"  [warn] fitz PDF conversion failed for {pdf_path}: {exc}")

    # --- attempt 2: pdf2image (pip install pdf2image; needs poppler) ---
    try:
        from pdf2image import convert_from_path  # type: ignore
        images = convert_from_path(str(pdf_path), first_page=1, last_page=1, dpi=144)
        if images:
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            images[0].save(tmp.name, "PNG")
            tmp.close()
            _g_temp_files.append(tmp.name)
            print(f"  [pdf→png] {pdf_path.name} -> {tmp.name} (pdf2image)")
            return Path(tmp.name)
    except ImportError:
        pass
    except Exception as exc:
        print(f"  [warn] pdf2image conversion failed for {pdf_path}: {exc}")

    print(f"  [warn] No PDF-to-image library found; cannot embed {pdf_path.name}.")
    print("         Install PyMuPDF:  pip install pymupdf")
    print("         or pdf2image:     pip install pdf2image")
    return None


def _resolve_image(figures_dir: str, rel: str) -> Optional[Path]:
    """
    Find an image file by trying the given path and common image extensions.
    PDF files are transparently converted to PNG before returning.
    """
    base = Path(figures_dir) / rel
    for candidate in [
        base,
        *[base.with_suffix(s) for s in (".png", ".jpg", ".jpeg", ".pdf")],
    ]:
        if candidate.exists():
            if candidate.suffix.lower() == ".pdf":
                return _pdf_to_image(candidate)   # convert PDF → PNG
            return candidate
    return None


def render_figure(doc, para: dict, figures_dir: str = ".") -> None:
    """
    Render a figure chunk.

    Detects minipage / subfigure layouts and uses a borderless Word table so
    that images and per-image captions stay side-by-side and properly aligned.
    Falls back to single-image layout otherwise.
    """
    translation = para.get("translation", para.get("text", ""))
    env_label   = para.get("env_label", "")

    # ── Side-by-side layout ──────────────────────────────────────────────
    if re.search(r'\\begin\{(?:minipage|subfigure|subfloat)\}', translation):
        minipages = _find_minipages(translation)
        if minipages:
            # Removed: doc.add_paragraph() blank line before table
            n   = len(minipages)
            tbl = doc.add_table(rows=2, cols=n)
            _remove_table_borders(tbl)
            col_w = Inches(6.0 / n * 0.9)

            for col, mp in enumerate(minipages):
                img_para = tbl.cell(0, col).paragraphs[0]
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                embedded = False
                for rel in mp["images"]:
                    found = _resolve_image(figures_dir, rel)
                    if found:
                        try:
                            img_para.add_run().add_picture(str(found), width=col_w)
                            embedded = True
                            break
                        except Exception as exc:
                            print(f"  [warn] subfigure {found}: {exc}")
                if not embedded and mp["images"]:
                    set_run_font(
                        img_para.add_run(f"[图片: {mp['images'][0]}]"),
                        FONT_BODY, 9, color=RGBColor(0x88, 0x88, 0x88),
                    )
                if mp["caption"]:
                    cap_para = tbl.cell(1, col).paragraphs[0]
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_run_font(
                        cap_para.add_run(_clean_spaces(mp["caption"])),
                        FONT_CAPTION, SIZE_CAPTION,
                    )

            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(lp, before_pt=4, after_pt=8)
            if env_label:
                set_run_font(lp.add_run(env_label), FONT_CAPTION, SIZE_CAPTION, bold=True)
            return

    # ── Single-image layout ──────────────────────────────────────────────
    img_paths     = re.findall(r'\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}', translation)
    caption_latex = _extract_caption(translation)
    # Removed: doc.add_paragraph() blank line before image

    embedded = False
    for rel in img_paths:
        found = _resolve_image(figures_dir, rel)
        if found:
            try:
                ip = doc.add_paragraph()
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ip.add_run().add_picture(str(found), width=Inches(5.0))
                embedded = True
            except Exception as exc:
                print(f"  [warn] figure {found}: {exc}")

    if not embedded and img_paths:
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade(ph, "F2F2F2")
        set_run_font(
            ph.add_run(f"[图片: {', '.join(img_paths)}]"),
            FONT_BODY, 10, color=RGBColor(0x88, 0x88, 0x88),
        )

    _add_caption(doc, env_label, caption_latex)
    # Removed: doc.add_paragraph() blank line after caption


def _set_table_auto_col_widths(doc, total_cm: float = 15.5) -> None:
    """
    Redistribute column widths of the most recently added table.

    Column widths are assigned proportionally to the longest text found in
    each column (measured in Unicode code-point count).  A minimum of 10 %
    of the total width is guaranteed per column so narrow columns are not
    squashed.  Falls back to equal widths when all columns are empty.

    total_cm should match the document's usable body width
    (page width − left margin − right margin; default matches the 15.5 cm
    body set up in main()).
    """
    if not doc.tables:
        return
    tbl    = doc.tables[-1]
    n_cols = len(tbl.columns)
    if n_cols == 0:
        return

    # --- measure max text length per column ---
    col_chars = [0] * n_cols
    for row in tbl.rows:
        cells = row.cells
        for i in range(min(n_cols, len(cells))):
            length = sum(len(p.text) for p in cells[i].paragraphs)
            if length > col_chars[i]:
                col_chars[i] = length

    total_chars = sum(col_chars)
    total_twips = int(total_cm * 567)   # 1 cm = 567 twips
    min_twips   = total_twips // (n_cols * 10)   # 10 % floor per column

    if total_chars == 0:
        col_twips = [total_twips // n_cols] * n_cols
    else:
        raw       = [int(total_twips * c / total_chars) for c in col_chars]
        col_twips = [max(w, min_twips) for w in raw]
        # Re-scale so the total stays exactly right
        excess = sum(col_twips) - total_twips
        if excess != 0:
            col_twips[col_twips.index(max(col_twips))] -= excess

    # --- apply to tblPr / tblW ---
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)

    for old_tag in (qn("w:tblW"), qn("w:tblLayout")):
        old = tblPr.find(old_tag)
        if old is not None:
            tblPr.remove(old)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    # --- rebuild tblGrid ---
    old_grid = tbl._tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl._tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for w in col_twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    tblPr_idx = list(tbl._tbl).index(tblPr)
    tbl._tbl.insert(tblPr_idx + 1, grid)

    # --- update every cell's tcW ---
    for row in tbl.rows:
        cells = row.cells
        for i in range(min(n_cols, len(cells))):
            tc    = cells[i]._tc
            tcPr  = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            old_tcW = tcPr.find(qn("w:tcW"))
            if old_tcW is not None:
                tcPr.remove(old_tcW)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_twips[i]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def render_table(doc, para: dict) -> None:
    """Render a table chunk via latex_table_renderer.
    Falls back to a raw-LaTeX placeholder block on parse failure.
    """
    source    = para.get("translation") or para.get("text", "")
    env_label = para.get("env_label", "")

    try:
        from table import parse_latex_table
        has_content = bool(parse_latex_table(source).get("rows"))
    except Exception:
        has_content = False

    if has_content:
        try:
            caption_latex = _extract_caption(source)    # Extract caption from LaTeX source
            _add_caption(doc, env_label, caption_latex) # Render label above the table
            render_latex_table_to_docx(doc, source)
            _set_table_auto_col_widths(doc)             # Redistribute column widths by content
            # Removed: doc.add_paragraph() blank line after table
            return
        except Exception as exc:
            print(f"  [table] Renderer failed for {env_label}: {exc}")

    # Fallback: amber placeholder (original behaviour)
    header = doc.add_paragraph()
    _shade(header, BG_TABLE)
    _spacing(header, before_pt=6, after_pt=2)
    set_run_font(
        header.add_run(f"[表格 {env_label}]"),
        FONT_BODY, 10, bold=True, color=RGBColor(0x80, 0x60, 0x00),
    )
    snippet = source[:600] + ("…" if len(source) > 600 else "")
    for line in snippet.split("\n"):
        p = doc.add_paragraph()
        _indent(p, left_cm=0.5)
        _shade(p, BG_TABLE)
        _spacing(p, before_pt=0, after_pt=0)
        run = p.add_run(line)
        run.font.name = FONT_CODE
        run.font.size = Pt(8)
    caption_latex = _extract_caption(source)
    if caption_latex:
        _add_caption(doc, env_label, caption_latex)
    # Removed: doc.add_paragraph() blank line after fallback block


def render_code(doc, para: dict) -> None:
    """Render verbatim / lstlisting as a grey monospace block (never translated)."""
    text      = para.get("text", "")
    env_label = para.get("env_label", "")

    body = text.strip()
    for env in ("verbatim", "lstlisting", "minted", "Verbatim"):
        _, inner = _extract_env_body(body, env)
        if inner:
            body = inner
            break
    else:
        body = re.sub(r'^\\begin\{[^}]+\}\*?', '', body)
        body = re.sub(r'\\end\{[^}]+\}\*?$',   '', body).strip()

    if env_label:
        lp = doc.add_paragraph()
        _spacing(lp, before_pt=8, after_pt=2)
        set_run_font(lp.add_run(env_label), FONT_BODY, SIZE_CODE, bold=True)

    for line in body.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        _indent(p, left_cm=0.5)
        _shade(p, BG_CODE)
        run = p.add_run(line)
        run.font.name = FONT_CODE
        run.font.size = Pt(SIZE_CODE)
    # Removed: doc.add_paragraph() blank line after code block


def render_mathenv(doc, para: dict) -> None:
    """Render theorem-like environments with a shaded header and indented body."""
    translation = para.get("translation", para.get("text", "")).strip()
    m           = _RE_MATHENV.match(translation)
    if not m:
        render_paragraph(doc, para)
        return

    env_name = m.group(1).lower().rstrip("*")
    label_zh = MATH_ENV_LABELS.get(env_name, env_name)
    opt, body = _extract_env_body(translation, env_name)
    if not body:
        render_paragraph(doc, para)
        return

    display_label = f"【{label_zh}】" + (f"（{opt}）" if opt else "")

    header = doc.add_paragraph()
    _shade(header, BG_THEOREM)
    _indent(header, left_cm=0.8)
    _spacing(header, before_pt=8, after_pt=0)
    set_run_font(header.add_run(display_label), FONT_BODY, SIZE_BODY, bold=True)

    ok = _inject_pandoc_paras(
        _clean_spaces(body), doc,
        font_name=FONT_BODY, size_pt=SIZE_BODY,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        indent_left_cm=0.8, before_pt=2.0, after_pt=8.0,
    )
    if not ok:
        fp = doc.add_paragraph()
        _shade(fp, BG_THEOREM)
        _indent(fp, left_cm=0.8)
        _spacing(fp, before_pt=2, after_pt=8)
        set_run_font(fp.add_run(_clean_spaces(body)), FONT_BODY, SIZE_BODY, italic=True)


def render_proof(doc, para: dict) -> None:
    translation = para.get("translation", para.get("text", "")).strip()
    opt, body   = _extract_env_body(translation, "proof")
    label_zh    = "证明" + (f"（{opt}）" if opt else "")

    lp = doc.add_paragraph()
    _indent(lp, left_cm=0.8)
    _spacing(lp, before_pt=4, after_pt=0)
    set_run_font(lp.add_run(label_zh + "  "), FONT_BODY, SIZE_BODY, bold=True)

    ok = _inject_pandoc_paras(
        _clean_spaces(body) if body else "", doc,
        font_name=FONT_BODY, size_pt=SIZE_BODY,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        indent_left_cm=0.8, before_pt=2.0, after_pt=4.0,
    )
    if not ok and body:
        fp = doc.add_paragraph()
        _indent(fp, left_cm=0.8)
        _spacing(fp, before_pt=2, after_pt=4)
        set_run_font(fp.add_run(_clean_spaces(body)), FONT_BODY, SIZE_BODY, italic=True)

    qp = doc.add_paragraph()
    qp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _spacing(qp, before_pt=0, after_pt=4)
    set_run_font(qp.add_run("□"), FONT_BODY, SIZE_BODY)


# ─────────────────────────────────────────────────────────────────────────────
# Reference list
# ─────────────────────────────────────────────────────────────────────────────

def render_references(doc, citations_path: str) -> None:
    """
    Append a reference list on a new page from citations.json (bib.py output).

    Schema: { "citekey": {"id": N, "citation": "..."}, ... }
    Entries are deduplicated by numeric id and sorted ascending.
    """
    path = Path(citations_path)
    if not path.exists():
        print(f"  [warn] citations file not found: {citations_path}")
        return
    try:
        with open(path, encoding="utf-8") as f:
            raw: Dict[str, Any] = json.load(f)
    except Exception as exc:
        print(f"  [warn] Could not load citations.json: {exc}")
        return

    # Deduplicate: first occurrence of each numeric id wins
    id_to_citation: Dict[int, str] = {}
    for entry in raw.values():
        try:
            nid = int(entry["id"])
            if nid not in id_to_citation:
                id_to_citation[nid] = str(entry["citation"])
        except (KeyError, TypeError, ValueError):
            pass

    if not id_to_citation:
        print("  [warn] citations.json has no valid entries -- skipping.")
        return

    doc.add_page_break()

    h = doc.add_heading("参考文献", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        set_run_font(run, FONT_HEADING, SIZE_H1, bold=True)
    _spacing(h, before_pt=0, after_pt=10)

    size = SIZE_BODY - 1   # 11 pt
    for nid in sorted(id_to_citation):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _indent(p, left_cm=1.2, hanging_cm=1.2)
        _spacing(p, before_pt=0, after_pt=4)
        _set_para_font(p, FONT_BODY, size)
        set_run_font(p.add_run(f"[{nid}] "), FONT_BODY, size, bold=True)
        set_run_font(p.add_run(_clean_spaces(id_to_citation[nid])), FONT_BODY, size)

    print(f"  [references] Appended {len(id_to_citation)} entries.")


# ─────────────────────────────────────────────────────────────────────────────
# Main dispatcher
# ─────────────────────────────────────────────────────────────────────────────

def render_chunk(doc, para: dict, figures_dir: str) -> None:
    """Classify and dispatch a single JSON paragraph to its renderer."""
    if para.get("syntax_error", False):
        print(f"  [skip] id={para.get('id')} (syntax_error flag)")
        return

    chunk_type = classify(para)
    print(f"  [{para.get('id', '?'):>4}] {para.get('env_label', ''):<22} -> {chunk_type}")

    {
        "abstract":      lambda: render_abstract(doc, para),
        "section":       lambda: render_heading(doc, para, 1),
        "subsection":    lambda: render_heading(doc, para, 2),
        "subsubsection": lambda: render_heading(doc, para, 3),
        "equation":      lambda: render_equation(doc, para),
        "figure":        lambda: render_figure(doc, para, figures_dir),
        "table":         lambda: render_table(doc, para),
        "code":          lambda: render_code(doc, para),
        "mathenv":       lambda: render_mathenv(doc, para),
        "proof":         lambda: render_proof(doc, para),
        "paragraph":     lambda: render_paragraph(doc, para),
    }.get(chunk_type, lambda: render_paragraph(doc, para))()


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    ap = argparse.ArgumentParser(
        description="Render replaced.json translation into a Word document."
    )
    ap.add_argument("--json",        default="replaced.json",
                    help="Input replaced.json (default: replaced.json)")
    ap.add_argument("--docx",        default="output.docx",
                    help="Output .docx -- created if absent, appended to if present")
    ap.add_argument("--figures-dir", default=".",
                    help="Base directory for resolving \\includegraphics paths")
    ap.add_argument("--citations",   default=None,
                    help="Path to citations.json -- appends a reference list page")
    ap.add_argument("--skip-title",  action="store_true",
                    help="Do not prepend the chapter title heading")
    args = ap.parse_args()

    json_path = Path(args.json)
    if not json_path.exists():
        sys.exit(f"Error: not found: {json_path}")
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)
    documents = data.get("documents", [])
    if not documents:
        sys.exit("Error: no documents found in JSON.")

    docx_path = Path(args.docx)
    if docx_path.exists():
        print(f"Opening : {docx_path}")
        doc = Document(str(docx_path))
    else:
        print(f"Creating: {docx_path}")
        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Cm(2.5)
            sec.bottom_margin = Cm(2.5)
            sec.left_margin   = Cm(3.0)
            sec.right_margin  = Cm(2.5)

    _init_footnotes(doc)

    for doc_meta in documents:
        paragraphs  = doc_meta.get("paragraphs", [])
        tex_file    = doc_meta.get("tex", "")
        figures_dir = str(Path(tex_file).parent) if tex_file else args.figures_dir

        print(f"\nDocument : {tex_file or '(unknown)'}")
        print(f"Title    : {doc_meta.get('title_translation', '')}")
        print(f"Chapter  : {doc_meta.get('chapter', '')}")
        print(f"Chunks   : {len(paragraphs)}")
        print(f"FigDir   : {figures_dir}\n")

        if not args.skip_title:
            render_chapter_title(doc, doc_meta)
        for para in paragraphs:
            render_chunk(doc, para, figures_dir)

    if args.citations:
        print("\nAppending reference list ...")
        render_references(doc, args.citations)

    doc.save(str(docx_path))
    print(f"\n[done] Saved -> {docx_path}")

    # Clean up any temporary files created during PDF→image conversion
    for tmp in _g_temp_files:
        try:
            os.unlink(tmp)
        except OSError:
            pass


if __name__ == "__main__":
    main()