from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from latex2mathml.converter import convert as latex_to_mathml
from mathml2omml import convert as mathml_to_omml


def add_omml_run(paragraph, omml_xml: str):
    """
    把 OMML (Office Math XML) 插入到 python-docx 的 paragraph 里。
    """
    # 有些转换器产出的 OMML 可能缺少命名空间声明；这里保险起见补一下
    # 典型 OMML 根是 <m:oMath> 或 <m:oMathPara>
    if "xmlns:m=" not in omml_xml:
        # 给根节点追加命名空间（尽量不破坏原结构）
        omml_xml = omml_xml.replace(
            "<m:oMath",
            f"<m:oMath {nsdecls('m')}",
            1
        ).replace(
            "<m:oMathPara",
            f"<m:oMathPara {nsdecls('m')}",
            1
        )

    omml_element = parse_xml(omml_xml)
    paragraph._p.append(omml_element)


def latex_to_omml(latex: str) -> str:
    """
    latex -> mathml -> omml
    返回 OMML 的 XML 字符串。
    """
    mathml = latex_to_mathml(latex)          # MathML string
    omml = mathml_to_omml(mathml)            # OMML string (XML)
    return omml


def add_math_inline(paragraph, latex: str, prefix_text: str = "", suffix_text: str = ""):
    """
    在段落里插入 inline 文本 + 公式 + 文本
    """
    if prefix_text:
        paragraph.add_run(prefix_text)
    omml = latex_to_omml(latex)
    add_omml_run(paragraph, omml)
    if suffix_text:
        paragraph.add_run(suffix_text)


def add_math_display(doc: Document, latex: str, title: str = None):
    """
    插入“独立一行”的 display 公式（简单做法：单独段落 + OMML）
    """
    if title:
        doc.add_paragraph(title)
    p = doc.add_paragraph()
    omml = latex_to_omml(latex)

    # 有些情况下 display 更像 <m:oMathPara>；mathml2omml 产出不一定
    # 这里直接插入即可，Word 通常能识别
    add_omml_run(p, omml)


if __name__ == "__main__":
    doc = Document()

    p1 = doc.add_paragraph("Inline demo: ")
    add_math_inline(
        p1,
        r"\alpha^2+\beta^2=\gamma^2",
        prefix_text="Pythagoras-like: ",
        suffix_text=" (editable in Word)"
    )

    add_math_display(doc, r"\int_0^\infty e^{-x^2}dx=\frac{\sqrt{\pi}}{2}", title="Display demo:")

    doc.save("latex_formula_demo.docx")
    print("Saved: latex_formula_demo.docx")