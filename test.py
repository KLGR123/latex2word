from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def set_run_font(run, font_name="Times New Roman", size=12, italic=False, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold


def ensure_footnotes_part(document):
    """
    Ensure the .docx has a footnotes part.
    Returns the footnotes XML element.
    """
    part = document.part
    package = part.package

    # Try to find existing footnotes relationship
    for rel in part.rels.values():
        if rel.reltype == RT.FOOTNOTES:
            footnotes_part = rel.target_part
            return footnotes_part._element

    # Create a new footnotes part if missing
    from docx.opc.packuri import PackURI
    from docx.opc.part import XmlPart
    from docx.oxml import parse_xml

    footnotes_partname = PackURI("/word/footnotes.xml")
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"

    xml = (
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:footnote w:id="-1" w:type="separator">'
        '<w:p><w:r><w:separator/></w:r></w:p>'
        '</w:footnote>'
        '<w:footnote w:id="0" w:type="continuationSeparator">'
        '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
        '</w:footnote>'
        '</w:footnotes>'
    )
    element = parse_xml(xml)
    footnotes_part = XmlPart(footnotes_partname, content_type, element, package)
    part.relate_to(footnotes_part, RT.FOOTNOTES)

    return footnotes_part._element


def get_next_footnote_id(footnotes_elm):
    ids = []
    for fn in footnotes_elm.findall(qn("w:footnote")):
        fid = fn.get(qn("w:id"))
        if fid is not None:
            try:
                ids.append(int(fid))
            except ValueError:
                pass
    positive_ids = [i for i in ids if i > 0]
    return max(positive_ids, default=0) + 1


def add_footnote(document, paragraph, footnote_text):
    """
    Add a real Word footnote to a paragraph.
    """
    footnotes_elm = ensure_footnotes_part(document)
    footnote_id = get_next_footnote_id(footnotes_elm)

    # 1) Add footnote reference in body paragraph
    run = paragraph.add_run()
    r = run._element
    rPr = OxmlElement("w:rPr")

    vert_align = OxmlElement("w:vertAlign")
    vert_align.set(qn("w:val"), "superscript")
    rPr.append(vert_align)
    r.append(rPr)

    footnote_ref = OxmlElement("w:footnoteReference")
    footnote_ref.set(qn("w:id"), str(footnote_id))
    r.append(footnote_ref)

    # 2) Add actual footnote content in footnotes.xml
    footnote = OxmlElement("w:footnote")
    footnote.set(qn("w:id"), str(footnote_id))

    p = OxmlElement("w:p")
    r1 = OxmlElement("w:r")
    ref = OxmlElement("w:footnoteRef")
    r1.append(ref)
    p.append(r1)

    r2 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = " " + footnote_text
    r2.append(t)
    p.append(r2)

    footnote.append(p)
    footnotes_elm.append(footnote)


def build_paragraph(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5

    # Segment 1
    r = p.add_run("The training process is facilitated by the open-source training frameworks ")
    set_run_font(r)

    r = p.add_run("Llama-Factory")
    set_run_font(r)
    add_footnote(document, p, "https://github.com/hiyouga/LLaMA-Factory")

    r = p.add_run(" and ")
    set_run_font(r)

    r = p.add_run("Unsloth")
    set_run_font(r)
    add_footnote(document, p, "https://github.com/unslothai/unsloth")

    r = p.add_run(
        " for DPO training, which offer convenient deployment options. "
        "We conducted model training on four 48GB L20 GPUs, and Table "
    )
    set_run_font(r)

    # Table reference
    r = p.add_run("1")
    set_run_font(r)

    r = p.add_run(
        " provides an overview of the approximate training and testing durations. "
        "During the training process, we set the validation set ratio to "
    )
    set_run_font(r)

    r = p.add_run("15%")
    set_run_font(r)

    r = p.add_run(" and the number of epoch to ")
    set_run_font(r)

    r = p.add_run("1")
    set_run_font(r)

    r = p.add_run(". We utilized the Low Rank Adaptation (LoRA) technique, setting the rank ")
    set_run_font(r)

    r = p.add_run("r")
    set_run_font(r, italic=True)

    r = p.add_run(" = 64. We set the deviation parameter to ")
    set_run_font(r)

    r = p.add_run("β")
    set_run_font(r, italic=True)

    r = p.add_run(" = 0.95. The learning rate was set to 0.0001, and the warm-up step ratio was set to 0.1.")
    set_run_font(r)


def main():
    doc = Document()

    # 页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    build_paragraph(doc)

    output_path = "training_paragraph_with_footnotes.docx"
    doc.save(output_path)
    print(f"Saved to {output_path}")


if __name__ == "__main__":
    main()