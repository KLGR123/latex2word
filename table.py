#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
table.py

Convert LaTeX tabular environments to python-docx Table objects.

Supported features:
  - Basic cells with & column separator and \\\\ row separator
  - booktabs rules: \\toprule, \\midrule, \\bottomrule
  - Standard \\hline
  - \\cmidrule[...]{m-n}  (partial horizontal rules)
  - \\multicolumn{n}{spec}{content}
  - \\multirow{n}{*}{content}  (basic, no real row-spanning in docx)
  - Nested inline formatting inside cells (passed through as plain text)
  - wraptable / longtable outer wrappers (inner tabular extracted)
  - Column spec parsing: l c r p{} X (alignment)
  - Caption extraction from surrounding table/wraptable environment

Public API
----------
    from table import parse_latex_table

    result = parse_latex_table(latex_text, doc, translation_text=None)
    # Returns TableResult(table, caption, col_count) or None on failure
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table, _Cell
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# Page content width in EMU (A4, 2.5cm margins each side → 16cm)
PAGE_WIDTH_CM = 16.0

FONT_REGULAR = "宋体"
FONT_CAPTION = "楷体"

# Border thickness in 1/8 pt (docx unit)
THICK_BORDER = 12   # ~1.5 pt  (toprule / bottomrule)
THIN_BORDER  = 6    # ~0.75 pt (midrule / hline)

TABLE_FONT_PT = 10.5


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class CellSpec:
    """Parsed content for a single table cell."""
    text: str = ""
    colspan: int = 1
    rowspan: int = 1          # tracked but not implemented in docx (limitation)
    align: str = "c"          # l / c / r
    bold: bool = False

@dataclass
class RowSpec:
    """One row of cells plus rule flags."""
    cells: List[CellSpec] = field(default_factory=list)
    rule_before: str = "none"   # none / hline / toprule / midrule / bottomrule
    rule_after:  str = "none"
    cmidrule_before: List[Tuple[int,int]] = field(default_factory=list)  # [(start,end),...]
    is_header: bool = False     # rows above first \midrule are header rows

@dataclass
class TableResult:
    table: Table
    caption: str = ""
    col_count: int = 0


# ---------------------------------------------------------------------------
# Brace / bracket helpers (same as in other modules)
# ---------------------------------------------------------------------------

def _find_closing_brace(s: str, start: int) -> int:
    """Return index AFTER matching '}' for '{' at s[start]."""
    assert s[start] == "{"
    depth = 1
    i = start + 1
    n = len(s)
    while i < n and depth:
        c = s[i]
        if c == "\\":
            i += 2
            continue
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
        i += 1
    return i


def _find_closing_bracket(s: str, start: int) -> int:
    """Return index AFTER ']' for '[' at s[start] (non-nested)."""
    i = start + 1
    n = len(s)
    while i < n and s[i] != "]":
        i += 1
    return i + 1


def _read_brace_arg(s: str, i: int) -> Tuple[str, int]:
    """Skip whitespace, read {content}, return (content, new_i)."""
    while i < len(s) and s[i] in " \t\n":
        i += 1
    if i >= len(s) or s[i] != "{":
        return "", i
    end = _find_closing_brace(s, i)
    return s[i+1:end-1], end


def _read_opt_arg(s: str, i: int) -> Tuple[str, int]:
    """Read optional [...] if present, else return ('', i)."""
    while i < len(s) and s[i] in " \t\n":
        i += 1
    if i >= len(s) or s[i] != "[":
        return "", i
    end = _find_closing_bracket(s, i)
    return s[i+1:end-1], end


# ---------------------------------------------------------------------------
# LaTeX comment / formatting strip (for cell content)
# ---------------------------------------------------------------------------

_COMMENT_RE = re.compile(r"(?<!\\)%[^\n]*")
_STRIP_CMDS = re.compile(
    r"\\(?:textbf|textit|emph|texttt|textrm|textsf|underline|"
    r"mathrm|mathbf|mathit|mathcal|mathbb|boldsymbol)\s*\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}",
    re.DOTALL,
)
_SPACING_RE = re.compile(
    r"\\(?:hspace\*?|vspace\*?|quad|qquad|,|;|!|:)\s*(?:\{[^}]*\})?"
)
_MATH_DISPLAY_RE = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
_MATH_INLINE_RE  = re.compile(r"\$(.+?)\$",    re.DOTALL)


def _clean_cell_text(raw: str) -> str:
    """Strip LaTeX formatting from cell text; preserve math $...$."""
    s = _COMMENT_RE.sub("", raw)
    # Remove spacing macros
    s = _SPACING_RE.sub(" ", s)
    # Unwrap formatting wrappers
    prev = None
    while prev != s:
        prev = s
        s = _STRIP_CMDS.sub(r"\1", s)
    # Normalize whitespace
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n+", " ", s)
    # Unescape common chars
    s = s.replace(r"\%", "%").replace(r"\&", "&").replace(r"\_", "_")
    s = s.replace(r"\{", "{").replace(r"\}", "}")
    s = s.replace(r"\textbackslash", "\\")
    s = s.replace("~", " ")
    s = s.strip()
    return s


def _is_bold_cell(raw: str) -> bool:
    """Return True if cell content is wrapped in \\textbf{...}."""
    s = raw.strip()
    return bool(re.match(r"^\\textbf\s*\{", s))


# ---------------------------------------------------------------------------
# Column spec parser
# ---------------------------------------------------------------------------

def parse_col_spec(spec: str) -> List[str]:
    """
    Parse a LaTeX column spec like 'lcr|p{3cm}X|c' into a list of
    alignment chars ('l', 'c', 'r') — one entry per data column,
    ignoring | separators and complex specs.
    """
    cols: List[str] = []
    i = 0
    n = len(spec)
    while i < n:
        c = spec[i]
        if c in "lcr":
            cols.append(c)
            i += 1
        elif c in "pXbm":
            cols.append("l")
            # skip optional args and brace arg if any
            i += 1
            if i < n and spec[i] == "{":
                end = _find_closing_brace(spec, i)
                i = end
        elif c == "@":
            i += 1
            if i < n and spec[i] == "{":
                end = _find_closing_brace(spec, i)
                i = end
        elif c == "*":
            # *{n}{spec}
            i += 1
            n_str, i = _read_brace_arg(spec, i)
            inner_spec, i = _read_brace_arg(spec, i)
            try:
                repeat = int(n_str.strip())
            except ValueError:
                repeat = 1
            inner = parse_col_spec(inner_spec)
            cols.extend(inner * repeat)
        else:
            i += 1  # skip | > < etc.
    return cols if cols else ["c"]


# ---------------------------------------------------------------------------
# Tabular body parser
# ---------------------------------------------------------------------------

_RULE_RE = re.compile(
    r"\\(toprule|midrule|bottomrule|hline|cline|cmidrule)\b"
)
_MULTICOLUMN_RE = re.compile(r"\\multicolumn\s*\{(\d+)\}\s*\{([^}]*)\}\s*\{")
_MULTIROW_RE    = re.compile(r"\\multirow\s*\{(\d+)\}\s*\{[^}]*\}\s*\{")


def _parse_multicolumn(s: str, m: re.Match) -> Tuple[CellSpec, int]:
    """Parse \\multicolumn{n}{spec}{content}, return (CellSpec, end_pos)."""
    colspan = int(m.group(1))
    spec_inner = m.group(2)
    aligns = [c for c in spec_inner if c in "lcr"]
    align = aligns[0] if aligns else "c"
    # Read content brace (m.end() points to '{' we already matched in group 3)
    content_start = m.end() - 1  # back to the opening '{'
    end = _find_closing_brace(s, content_start)
    content = s[content_start+1:end-1]
    is_bold = _is_bold_cell(content)
    return CellSpec(
        text=_clean_cell_text(content),
        colspan=colspan,
        align=align,
        bold=is_bold,
    ), end


def _parse_cell_content(s: str, col_align: str) -> CellSpec:
    """Build CellSpec from raw cell string."""
    is_bold = _is_bold_cell(s)
    return CellSpec(
        text=_clean_cell_text(s),
        align=col_align,
        bold=is_bold,
    )


def parse_tabular_body(
    body: str,
    col_specs: List[str],
) -> List[RowSpec]:
    """
    Parse the content inside \\begin{tabular}...\\end{tabular}.

    Returns a list of RowSpec, one per data row.
    """
    rows: List[RowSpec] = []
    # Split into logical lines by \\  (row separator)
    # We need to be careful: \\\\ inside math should not split
    # Simple approach: split on \\\\ not inside $...$
    logical_lines = _split_rows(body)

    current_rule_before = "none"
    cmidrule_before: List[Tuple[int,int]] = []

    for raw_line in logical_lines:
        line = raw_line.strip()
        if not line:
            continue

        # Check if line is purely rules
        rule_tokens = _extract_rule_tokens(line)
        if rule_tokens is not None:
            # Pure rule line(s)
            for tok in rule_tokens:
                if tok[0] == "cmidrule":
                    cmidrule_before.append(tok[1])
                else:
                    current_rule_before = tok[0]
            continue

        # Data row — may start with some rule tokens
        rule_prefix, data_part = _split_rule_prefix(line)
        for tok in rule_prefix:
            if tok[0] == "cmidrule":
                cmidrule_before.append(tok[1])
            else:
                current_rule_before = tok[0]

        if not data_part.strip():
            continue

        cells = _parse_row_cells(data_part, col_specs)
        if not cells:
            continue

        row = RowSpec(
            cells=cells,
            rule_before=current_rule_before,
            cmidrule_before=cmidrule_before,
        )
        rows.append(row)
        current_rule_before = "none"
        cmidrule_before = []

    # Mark header rows: everything before (and including) first midrule row
    _mark_header_rows(rows)
    return rows


def _split_rows(body: str) -> List[str]:
    """Split tabular body on \\\\ respecting $...$ math zones."""
    result: List[str] = []
    i = 0
    n = len(body)
    start = 0
    in_math = False

    while i < n:
        c = body[i]
        if c == "$":
            # Toggle math mode (simple: doesn't handle $$)
            in_math = not in_math
            i += 1
        elif c == "\\" and not in_math:
            if i + 1 < n and body[i+1] == "\\":
                result.append(body[start:i])
                i += 2
                # Skip optional [dimspec] after \\
                while i < n and body[i] in " \t":
                    i += 1
                if i < n and body[i] == "[":
                    end = _find_closing_bracket(body, i)
                    i = end
                start = i
            else:
                i += 1
        else:
            i += 1
    tail = body[start:].strip()
    if tail:
        result.append(tail)
    return result


def _extract_rule_tokens(line: str) -> Optional[List]:
    """
    If line consists ONLY of rule commands, return list of token tuples.
    Otherwise return None.
    """
    remainder = line.strip()
    tokens = []
    while remainder:
        m = re.match(r"\\(toprule|midrule|bottomrule|hline)\b(\s*)", remainder)
        if m:
            tokens.append((m.group(1), None))
            remainder = remainder[m.end():].strip()
            continue
        m = re.match(r"\\cmidrule(?:\s*\[[^\]]*\])?\s*\{(\d+)-(\d+)\}", remainder)
        if m:
            tokens.append(("cmidrule", (int(m.group(1)), int(m.group(2)))))
            remainder = remainder[m.end():].strip()
            continue
        m = re.match(r"\\cline\s*\{(\d+)-(\d+)\}", remainder)
        if m:
            tokens.append(("cmidrule", (int(m.group(1)), int(m.group(2)))))
            remainder = remainder[m.end():].strip()
            continue
        # Not a pure rule line
        return None
    return tokens if tokens else None


def _split_rule_prefix(line: str) -> Tuple[List, str]:
    """
    Strip any leading rule commands from a data line.
    Returns (rule_tokens, remaining_data).
    """
    tokens = []
    s = line.strip()
    while True:
        m = re.match(r"\\(toprule|midrule|bottomrule|hline)\b\s*", s)
        if m:
            tokens.append((m.group(1), None))
            s = s[m.end():].strip()
            continue
        m = re.match(r"\\cmidrule(?:\s*\[[^\]]*\])?\s*\{(\d+)-(\d+)\}\s*", s)
        if m:
            tokens.append(("cmidrule", (int(m.group(1)), int(m.group(2)))))
            s = s[m.end():].strip()
            continue
        break
    return tokens, s


def _parse_row_cells(row_text: str, col_specs: List[str]) -> List[CellSpec]:
    """Split a row on '&' and build CellSpec list."""
    # Split on & not inside braces
    parts = _split_on_ampersand(row_text)
    cells: List[CellSpec] = []
    col_idx = 0

    for raw in parts:
        raw = raw.strip()
        if not raw:
            # Empty cell
            align = col_specs[col_idx] if col_idx < len(col_specs) else "l"
            cells.append(CellSpec(text="", align=align))
            col_idx += 1
            continue

        # Check for \multicolumn
        mc = _MULTICOLUMN_RE.match(raw)
        if mc:
            cell, _ = _parse_multicolumn(raw, mc)
            cells.append(cell)
            col_idx += cell.colspan
            continue

        # Check for \multirow (basic support: extract content only)
        mr = _MULTIROW_RE.match(raw)
        if mr:
            content_start = mr.end() - 1
            end = _find_closing_brace(raw, content_start)
            content = raw[content_start+1:end-1]
            align = col_specs[col_idx] if col_idx < len(col_specs) else "l"
            cells.append(CellSpec(
                text=_clean_cell_text(content),
                align=align,
                bold=_is_bold_cell(content),
            ))
            col_idx += 1
            continue

        # Normal cell
        align = col_specs[col_idx] if col_idx < len(col_specs) else "l"
        cells.append(_parse_cell_content(raw, align))
        col_idx += 1

    return cells


def _split_on_ampersand(s: str) -> List[str]:
    """Split string on '&' not inside braces or $...$."""
    parts: List[str] = []
    depth = 0
    in_math = False
    start = 0
    i = 0
    n = len(s)
    while i < n:
        c = s[i]
        if c == "$":
            in_math = not in_math
        elif c == "\\" and i + 1 < n:
            i += 2
            continue
        elif not in_math:
            if c == "{":
                depth += 1
            elif c == "}":
                depth -= 1
            elif c == "&" and depth == 0:
                parts.append(s[start:i])
                start = i + 1
        i += 1
    parts.append(s[start:])
    return parts


def _mark_header_rows(rows: List[RowSpec]) -> None:
    r"""
    Rows before (and including) the first midrule/hline are header rows.
    For booktabs tables: rows before first \midrule.
    For hline tables: first row after \toprule or \hline.
    """
    found_mid = False
    for row in rows:
        if found_mid:
            break
        if row.rule_before in ("toprule",):
            row.is_header = True
            continue
        if row.rule_before in ("midrule", "hline"):
            found_mid = True
            break
        if not found_mid:
            row.is_header = True


# ---------------------------------------------------------------------------
# Caption extractor
# ---------------------------------------------------------------------------

_CAPTION_RE = re.compile(
    r"\\caption\*?\s*(?:\[[^\]]*\])?\s*\{", re.DOTALL
)
_LABEL_RE = re.compile(r"\\label\s*\{[^}]*\}")


def extract_caption(latex_text: str) -> str:
    """Extract and clean first \\caption{...} from a table environment."""
    m = _CAPTION_RE.search(latex_text)
    if not m:
        return ""
    start = m.end() - 1   # position of '{'
    end = _find_closing_brace(latex_text, start)
    raw = latex_text[start+1:end-1]
    # Remove \label inside caption
    raw = _LABEL_RE.sub("", raw)
    return _clean_cell_text(raw)


def extract_caption_from_translation(translation: str) -> str:
    """Same as extract_caption but prefer the translation field."""
    return extract_caption(translation)


# ---------------------------------------------------------------------------
# Tabular content extractor
# ---------------------------------------------------------------------------

_BEGIN_TABULAR_RE = re.compile(
    r"\\begin\{(tabular\*?|tabularx|tabulary|tabularray|longtable|array)\}"
)
_END_TABULAR_RE = re.compile(
    r"\\end\{(tabular\*?|tabularx|tabulary|tabularray|longtable|array)\}"
)


def extract_tabular(latex_text: str) -> Optional[Tuple[str, str]]:
    """
    Find the innermost tabular-like environment.
    Returns (col_spec, body) or None.
    """
    m_begin = _BEGIN_TABULAR_RE.search(latex_text)
    if not m_begin:
        return None

    env_name = m_begin.group(1)
    after = latex_text[m_begin.end():]

    # tabularx has extra {width} arg before col spec; skip it
    if env_name in ("tabularx", "tabulary"):
        if after.lstrip().startswith("{"):
            _, skip_end = _read_brace_arg(after, after.index("{"))
            after = after[skip_end:]

    # Read optional [] (tblr style)
    _, after_stripped = _read_opt_arg(after, 0)
    if after_stripped:
        after = after[len(after) - len(after_stripped):]

    # Read col spec
    if not after.lstrip().startswith("{"):
        # No brace for col spec — try tabularray = {...} style
        return None
    brace_idx  = after.index("{")
    col_spec, end_idx = _read_brace_arg(after, brace_idx)
    rest = after[end_idx:]   # everything after the col-spec brace

    # Body is everything until matching \end{env}
    end_pattern = re.compile(r"\\end\{" + re.escape(env_name) + r"\}")
    m_end = end_pattern.search(rest)
    if not m_end:
        return None

    body = rest[:m_end.start()]
    return col_spec, body


# ---------------------------------------------------------------------------
# python-docx table builder
# ---------------------------------------------------------------------------

_ALIGN_MAP = {
    "l": WD_ALIGN_PARAGRAPH.LEFT,
    "c": WD_ALIGN_PARAGRAPH.CENTER,
    "r": WD_ALIGN_PARAGRAPH.RIGHT,
}

def _pt_to_emu(pt: float) -> int:
    return int(pt * 12700)

def _set_cell_border(cell: _Cell, sides: dict) -> None:
    """
    Set cell borders.
    sides = {'top': thick, 'bottom': thick, ...}  thickness in 1/8 pt units
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, size in sides.items():
        el = OxmlElement(f"w:{side}")
        if size == 0:
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _no_border() -> dict:
    return {"top": 0, "bottom": 0, "left": 0, "right": 0}


def _add_cell_text(
    cell: _Cell,
    text: str,
    align: str = "c",
    bold: bool = False,
    is_header: bool = False,
    font_size_pt: float = TABLE_FONT_PT,
    font_name: str = FONT_REGULAR,
) -> None:
    """Write text into a docx cell with proper formatting."""
    para = cell.paragraphs[0]
    para.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    run.bold = bold or is_header
    run.font.size = Pt(font_size_pt)
    run.font.name = font_name
    # Set East Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def build_docx_table(
    doc: Document,
    rows: List[RowSpec],
    col_specs: List[str],
) -> Table:
    """
    Build a python-docx Table from parsed RowSpec list.
    Implements 三线表 (three-line table) style by default for booktabs.
    """
    if not rows:
        raise ValueError("No rows to build table from")

    # Determine actual column count (respecting multicolumn)
    n_cols = max(
        sum(c.colspan for c in row.cells)
        for row in rows
    ) if rows else len(col_specs)
    n_cols = max(n_cols, len(col_specs), 1)

    n_rows = len(rows)

    # Compute column widths (equal distribution)
    page_width_dxa = int(PAGE_WIDTH_CM / 2.54 * 1440)
    col_width_dxa  = page_width_dxa // n_cols

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Detect booktabs style
    has_booktabs = any(
        r.rule_before in ("toprule", "midrule", "bottomrule")
        for r in rows
    )

    for row_idx, row_spec in enumerate(rows):
        docx_row = table.rows[row_idx]
        col_offset = 0

        for cell_spec in row_spec.cells:
            if col_offset >= n_cols:
                break

            # Merge spanning cells
            if cell_spec.colspan > 1:
                end_col = min(col_offset + cell_spec.colspan - 1, n_cols - 1)
                cell = docx_row.cells[col_offset]
                merge_to = docx_row.cells[end_col]
                cell = cell.merge(merge_to)
            else:
                cell = docx_row.cells[col_offset]

            _add_cell_text(
                cell,
                cell_spec.text,
                align=cell_spec.align,
                bold=cell_spec.bold,
                is_header=row_spec.is_header,
            )
            col_offset += cell_spec.colspan

        # Fill any remaining cells (in case row had fewer cells than cols)
        while col_offset < n_cols:
            _add_cell_text(docx_row.cells[col_offset], "")
            col_offset += 1

    # Apply borders: 三线表 for booktabs, simple hline for others
    _apply_borders(table, rows, n_cols, has_booktabs)

    # Set column widths
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(col_width_dxa))
            tcW.set(qn("w:type"), "dxa")

    return table


def _apply_borders(
    table: Table,
    rows: List[RowSpec],
    n_cols: int,
    has_booktabs: bool,
) -> None:
    """
    Apply border rules to every cell based on RowSpec metadata.
    三线表: thick top/bottom, thin mid; no vertical borders.
    hline: thin horizontal borders, no vertical.
    """
    n_rows = len(rows)

    for row_idx, row_spec in enumerate(rows):
        is_last = row_idx == n_rows - 1

        for col_idx in range(n_cols):
            cell = table.rows[row_idx].cells[col_idx]

            top_size    = 0
            bottom_size = 0

            if has_booktabs:
                # Top border of this row
                if row_spec.rule_before == "toprule":
                    top_size = THICK_BORDER
                elif row_spec.rule_before in ("midrule",):
                    top_size = THIN_BORDER
                elif row_spec.rule_before == "hline":
                    top_size = THIN_BORDER

                # Bottom border: last row gets bottomrule (thick)
                if is_last:
                    # Check if following "rule_after" or just last row
                    bottom_size = THICK_BORDER
                elif row_idx + 1 < n_rows:
                    next_rule = rows[row_idx + 1].rule_before
                    if next_rule == "bottomrule":
                        bottom_size = THICK_BORDER
                    # midrule is handled as top of next row
            else:
                # hline style
                if row_spec.rule_before in ("hline", "toprule"):
                    top_size = THIN_BORDER
                if is_last:
                    bottom_size = THIN_BORDER
                elif row_idx + 1 < n_rows:
                    if rows[row_idx + 1].rule_before in ("hline", "midrule", "bottomrule"):
                        bottom_size = THIN_BORDER

            sides = {
                "top":    top_size,
                "bottom": bottom_size,
                "left":   0,
                "right":  0,
                "insideH": 0,
                "insideV": 0,
            }
            _set_cell_border(cell, sides)

        # Handle cmidrule (partial bottom border on previous row)
        if row_spec.cmidrule_before:
            prev_idx = row_idx - 1
            if prev_idx >= 0:
                for (start_col, end_col) in row_spec.cmidrule_before:
                    for c in range(start_col - 1, min(end_col, n_cols)):
                        prev_cell = table.rows[prev_idx].cells[c]
                        tc = prev_cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        # Add bottom border only
                        tcBorders = tcPr.find(qn("w:tcBorders"))
                        if tcBorders is None:
                            tcBorders = OxmlElement("w:tcBorders")
                            tcPr.append(tcBorders)
                        bot = OxmlElement("w:bottom")
                        bot.set(qn("w:val"), "single")
                        bot.set(qn("w:sz"), str(THIN_BORDER))
                        bot.set(qn("w:color"), "000000")
                        tcBorders.append(bot)


# ---------------------------------------------------------------------------
# tblr / tabularray basic parser
# ---------------------------------------------------------------------------

def extract_tblr_body(latex_text: str) -> Optional[Tuple[str, str]]:
    """
    Attempt to extract col spec and body from \\begin{tblr}[...]{spec}...\\end{tblr}.
    Very basic: treats the first {spec} as column spec.
    """
    m = re.search(r"\\begin\{(?:tblr|tabularray|longtblr)\}", latex_text)
    if not m:
        return None
    after = latex_text[m.end():]
    # Skip optional []
    _, after2 = _read_opt_arg(after, 0)
    remainder = after[len(after) - len(after2):] if after2 else after
    if not remainder.lstrip().startswith("{"):
        # Try body directly
        col_spec = "c"
        body_end_m = re.search(r"\\end\{(?:tblr|tabularray|longtblr)\}", latex_text[m.end():])
        if not body_end_m:
            return None
        body = latex_text[m.end():m.end() + body_end_m.start()]
        return col_spec, body

    col_spec, rest = _read_brace_arg(remainder, remainder.index("{"))
    # For tblr, col spec may contain = style; extract just alignment chars
    end_pattern = re.compile(r"\\end\{(?:tblr|tabularray|longtblr)\}")
    m_end = end_pattern.search(rest)
    if not m_end:
        return None
    body = rest[:m_end.start()]
    return col_spec, body


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def parse_latex_table(
    latex_text: str,
    doc: Document,
    translation_text: Optional[str] = None,
) -> Optional[TableResult]:
    """
    Parse a LaTeX table block and return a TableResult.

    Parameters
    ----------
    latex_text : str
        Raw LaTeX for the table (may include \\begin{table}...\\end{table} wrapper).
    doc : Document
        python-docx Document to add the table to.
    translation_text : str, optional
        Translation of the same block; used to extract the caption in Chinese.

    Returns
    -------
    TableResult or None
        None means parsing failed; caller should insert a placeholder.
    """
    # Extract caption from translation (prefer translated caption)
    caption = ""
    if translation_text:
        caption = extract_caption_from_translation(translation_text)
    if not caption:
        caption = extract_caption(latex_text)

    # Try standard tabular
    result = extract_tabular(latex_text)

    # Fallback: tblr / tabularray
    if result is None:
        result = extract_tblr_body(latex_text)

    if result is None:
        return None

    col_spec_str, body = result

    try:
        col_specs = parse_col_spec(col_spec_str)
        rows = parse_tabular_body(body, col_specs)
    except Exception as exc:
        print(f"[table] Parse error: {exc}", file=sys.stderr)
        return None

    if not rows:
        return None

    try:
        table = build_docx_table(doc, rows, col_specs)
    except Exception as exc:
        print(f"[table] Build error: {exc}", file=sys.stderr)
        return None

    n_cols = max(sum(c.colspan for c in r.cells) for r in rows)
    return TableResult(table=table, caption=caption, col_count=n_cols)