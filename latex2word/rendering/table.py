import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# LaTeX math symbol → Unicode mapping
# ---------------------------------------------------------------------------

MATH_SYMBOL_MAP = {
    "\\dagger":    "†",
    "\\ddagger":   "‡",
    "\\sharp":     "♯",
    "\\flat":      "♭",
    "\\natural":   "♮",
    "\\star":      "★",
    "\\ast":       "*",
    "\\bullet":    "•",
    "\\circ":      "○",
    "\\times":     "×",
    "\\div":       "÷",
    "\\pm":        "±",
    "\\mp":        "∓",
    "\\infty":     "∞",
    "\\leq":       "≤",
    "\\geq":       "≥",
    "\\neq":       "≠",
    "\\approx":    "≈",
    "\\sim":       "~",
    "\\equiv":     "≡",
    "\\rightarrow":"→",
    "\\leftarrow": "←",
    "\\Rightarrow":"⇒",
    "\\Leftarrow": "⇐",
    "\\alpha":     "α",
    "\\beta":      "β",
    "\\gamma":     "γ",
    "\\delta":     "δ",
    "\\epsilon":   "ε",
    "\\zeta":      "ζ",
    "\\eta":       "η",
    "\\theta":     "θ",
    "\\iota":      "ι",
    "\\kappa":     "κ",
    "\\lambda":    "λ",
    "\\mu":        "μ",
    "\\nu":        "ν",
    "\\xi":        "ξ",
    "\\pi":        "π",
    "\\rho":       "ρ",
    "\\sigma":     "σ",
    "\\tau":       "τ",
    "\\phi":       "φ",
    "\\psi":       "ψ",
    "\\omega":     "ω",
    "\\Gamma":     "Γ",
    "\\Delta":     "Δ",
    "\\Theta":     "Θ",
    "\\Lambda":    "Λ",
    "\\Pi":        "Π",
    "\\Sigma":     "Σ",
    "\\Phi":       "Φ",
    "\\Psi":       "Ψ",
    "\\Omega":     "Ω",
    "\\%":         "%",
    "\\&":         "&",
    "\\$":         "$",
    "\\_":         "_",
    "\\#":         "#",
    "\\{":         "{",
    "\\}":         "}",
    "\\textdagger":"†",
    "\\textpm":    "±",
}

# Commands that are silently consumed (no output)
_SKIP_CMDS = frozenset({
    "hline", "toprule", "midrule", "bottomrule", "cline",
    "noindent", "centering", "raggedright", "raggedleft",
    "small", "footnotesize", "scriptsize", "tiny",
    "normalsize", "large", "Large", "LARGE", "huge", "Huge",
    "vspace", "hspace", "vskip", "hskip",
    "label", "ref", "cite", "nocite",
    "arraybackslash", "newline", "par",
})


# ---------------------------------------------------------------------------
# Low-level LaTeX text utilities
# ---------------------------------------------------------------------------

def _extract_braced(s: str, start: int):
    """
    Extract the content of a brace-balanced group starting at s[start] == '{'.

    Returns (content: str, end: int) where end is the index after the closing '}'.
    Falls back gracefully if braces are unbalanced.
    """
    if start >= len(s) or s[start] != "{":
        return "", start
    depth = 0
    i = start
    while i < len(s):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return s[start + 1 : i], i + 1
        i += 1
    return s[start + 1 :], len(s)


def _extract_optional(s: str, start: int):
    """
    Extract optional argument [...] starting at s[start] == '['.

    Returns (content: str, end: int).
    """
    if start >= len(s) or s[start] != "[":
        return "", start
    end = s.find("]", start + 1)
    if end == -1:
        return s[start + 1 :], len(s)
    return s[start + 1 : end], end + 1


def _extract_math_dollar(s: str, start: int):
    """
    Extract content of $...$ starting at s[start] == '$'.

    Returns (content: str, end: int).  Handles $$ as empty math.
    """
    if start + 1 < len(s) and s[start + 1] == "$":
        return "", start + 2  # $$ display math — treat as empty
    end = s.find("$", start + 1)
    if end == -1:
        return s[start + 1 :], len(s)
    return s[start + 1 : end], end + 1


# ---------------------------------------------------------------------------
# Span model and parser
# ---------------------------------------------------------------------------

def _span(text, bold=False, italic=False, underline=False,
          superscript=False, subscript=False):
    return {
        "text":        text,
        "bold":        bold,
        "italic":      italic,
        "underline":   underline,
        "superscript": superscript,
        "subscript":   subscript,
    }


def _same_style(a, b):
    return (a["bold"] == b["bold"] and
            a["italic"] == b["italic"] and
            a["underline"] == b["underline"] and
            a["superscript"] == b["superscript"] and
            a["subscript"] == b["subscript"])


def _merge_spans(spans):
    """Merge consecutive spans that share the same style."""
    out = []
    for s in spans:
        if not s["text"]:
            continue
        if out and _same_style(out[-1], s):
            out[-1] = dict(out[-1], text=out[-1]["text"] + s["text"])
        else:
            out.append(s)
    return out


def _parse_segment(s: str, spans: list, *,
                   bold=False, italic=False, underline=False,
                   superscript=False, subscript=False):
    """
    Recursively scan a LaTeX fragment and append span dicts to *spans*.
    Inherits the current formatting context (bold, italic, …).
    """
    i = 0
    buf = ""

    def flush():
        nonlocal buf
        if buf:
            spans.append(_span(buf, bold=bold, italic=italic, underline=underline,
                                superscript=superscript, subscript=subscript))
            buf = ""

    def recurse(text, **overrides):
        """Recurse with updated formatting context."""
        ctx = dict(bold=bold, italic=italic, underline=underline,
                   superscript=superscript, subscript=subscript)
        ctx.update(overrides)
        _parse_segment(text, spans, **ctx)

    while i < len(s):
        c = s[i]

        # ----------------------------------------------------------------
        # Inline math  $...$
        # ----------------------------------------------------------------
        if c == "$":
            flush()
            math, i = _extract_math_dollar(s, i)
            math = math.strip()
            if math.startswith("^"):
                inner = math[1:]
                if inner.startswith("{"):
                    inner, _ = _extract_braced(inner, 0)
                inner = _resolve_math(inner)
                spans.append(_span(inner, bold=bold, italic=italic, underline=underline,
                                   superscript=True, subscript=False))
            elif math.startswith("_"):
                inner = math[1:]
                if inner.startswith("{"):
                    inner, _ = _extract_braced(inner, 0)
                inner = _resolve_math(inner)
                spans.append(_span(inner, bold=bold, italic=italic, underline=underline,
                                   superscript=False, subscript=True))
            else:
                # First try a pure symbol-map resolution (fast path for
                # single symbols like \\alpha, \\dagger, \\sharp …).
                # If the result still contains a backslash the content has
                # structure (e.g. \\hat{e}, \\frac{a}{b}) — fall back to
                # _parse_segment so unknown commands strip down to their
                # brace argument (best-effort plain-text rendering).
                resolved = _resolve_math(math)
                if "\\" in resolved:
                    recurse(math)  # re-parse as LaTeX text fragment
                elif resolved:
                    spans.append(_span(resolved, bold=bold, italic=italic,
                                       underline=underline,
                                       superscript=superscript,
                                       subscript=subscript))
            continue

        # ----------------------------------------------------------------
        # LaTeX command  \cmdname
        # ----------------------------------------------------------------
        if c == "\\":
            j = i + 1
            if j < len(s) and s[j].isalpha():
                # Read full alphabetic command name
                while j < len(s) and s[j].isalpha():
                    j += 1
                cmd = s[i + 1 : j]
                i = j
                # Skip trailing space after command name
                while i < len(s) and s[i] == " ":
                    i += 1

                # --- Formatting commands ---
                if cmd in ("textbf", "mathbf"):
                    flush()
                    if i < len(s) and s[i] == "{":
                        content, i = _extract_braced(s, i)
                        recurse(content, bold=True)
                    continue

                if cmd in ("textit", "mathit", "emph"):
                    flush()
                    if i < len(s) and s[i] == "{":
                        content, i = _extract_braced(s, i)
                        recurse(content, italic=True)
                    continue

                if cmd in ("underline", "uline", "uuline"):
                    flush()
                    if i < len(s) and s[i] == "{":
                        content, i = _extract_braced(s, i)
                        recurse(content, underline=True)
                    continue

                if cmd in ("textsuperscript", "textsu"):
                    flush()
                    if i < len(s) and s[i] == "{":
                        content, i = _extract_braced(s, i)
                        recurse(content, superscript=True, subscript=False)
                    continue

                if cmd in ("textsubscript", "textsb"):
                    flush()
                    if i < len(s) and s[i] == "{":
                        content, i = _extract_braced(s, i)
                        recurse(content, superscript=False, subscript=True)
                    continue

                # --- Font shape switches (scoped by following braced group or rest) ---
                if cmd == "bfseries":
                    flush()
                    if i < len(s) and s[i] == "{":
                        content, i = _extract_braced(s, i)
                        recurse(content, bold=True)
                    else:
                        bold = True  # affects rest of current scope
                    continue

                if cmd == "itshape":
                    flush()
                    if i < len(s) and s[i] == "{":
                        content, i = _extract_braced(s, i)
                        recurse(content, italic=True)
                    else:
                        italic = True
                    continue

                # --- Commands with mandatory args to consume silently ---
                if cmd in ("cellcolor", "rowcolor", "color", "colorbox",
                           "fcolorbox", "textcolor"):
                    # Consume optional [model] arg
                    if i < len(s) and s[i] == "[":
                        _, i = _extract_optional(s, i)
                    # Consume one or two mandatory {color} args
                    if i < len(s) and s[i] == "{":
                        _, i = _extract_braced(s, i)
                    if cmd == "fcolorbox" and i < len(s) and s[i] == "{":
                        _, i = _extract_braced(s, i)
                    # For \textcolor{color}{text}, render the text
                    if cmd == "textcolor" and i < len(s) and s[i] == "{":
                        flush()
                        content, i = _extract_braced(s, i)
                        recurse(content)
                    continue

                if cmd in ("multicolumn",):
                    # \multicolumn{n}{align}{content}
                    for _ in range(2):
                        if i < len(s) and s[i] == "{":
                            _, i = _extract_braced(s, i)
                    if i < len(s) and s[i] == "{":
                        flush()
                        content, i = _extract_braced(s, i)
                        recurse(content)
                    continue

                if cmd == "multirow":
                    # \multirow{n}{width}{content}
                    for _ in range(2):
                        if i < len(s) and s[i] == "{":
                            _, i = _extract_braced(s, i)
                    if i < len(s) and s[i] == "{":
                        flush()
                        content, i = _extract_braced(s, i)
                        recurse(content)
                    continue

                if cmd in ("text", "mbox", "hbox", "textrm", "texttt",
                           "textsf", "textsc", "textmd", "textnormal"):
                    flush()
                    if i < len(s) and s[i] == "{":
                        content, i = _extract_braced(s, i)
                        recurse(content)
                    continue

                # vspace / hspace: consume argument, no output
                if cmd in ("vspace", "hspace"):
                    if i < len(s) and s[i] == "{":
                        _, i = _extract_braced(s, i)
                    continue

                if cmd in _SKIP_CMDS:
                    # Some of these may have an argument we should consume
                    if cmd in ("vspace", "hspace", "cline", "label", "ref",
                               "cite", "nocite") and i < len(s) and s[i] == "{":
                        _, i = _extract_braced(s, i)
                    continue

                # Check math symbol map
                sym = MATH_SYMBOL_MAP.get("\\" + cmd)
                if sym:
                    flush()
                    spans.append(_span(sym, bold=bold, italic=italic,
                                       underline=underline,
                                       superscript=superscript,
                                       subscript=subscript))
                    continue

                # Unknown command with braced arg: render the arg content
                if i < len(s) and s[i] == "{":
                    flush()
                    content, i = _extract_braced(s, i)
                    recurse(content)
                # else: silently ignore unknown command

            else:
                # Single-character special command: \\, \%, \&, etc.
                if j < len(s):
                    sc = s[j]
                    sym = MATH_SYMBOL_MAP.get("\\" + sc)
                    if sym:
                        flush()
                        spans.append(_span(sym, bold=bold, italic=italic,
                                           underline=underline,
                                           superscript=superscript,
                                           subscript=subscript))
                    elif sc == "\\":
                        pass  # line break in table, skip
                    i = j + 1
                else:
                    i += 1
            continue

        # ----------------------------------------------------------------
        # Explicit braced group  {...}
        # ----------------------------------------------------------------
        if c == "{":
            flush()
            content, i = _extract_braced(s, i)
            recurse(content)
            continue

        if c == "}":
            i += 1
            continue

        # ----------------------------------------------------------------
        # Bare superscript / subscript  ^{} / _{}
        # ----------------------------------------------------------------
        if c == "^":
            flush()
            i += 1
            if i < len(s) and s[i] == "{":
                content, i = _extract_braced(s, i)
            elif i < len(s):
                content, i = s[i], i + 1
            else:
                content = ""
            content = _resolve_math(content)
            spans.append(_span(content, bold=bold, italic=italic,
                               underline=underline, superscript=True))
            continue

        if c == "_":
            flush()
            i += 1
            if i < len(s) and s[i] == "{":
                content, i = _extract_braced(s, i)
            elif i < len(s):
                content, i = s[i], i + 1
            else:
                content = ""
            content = _resolve_math(content)
            spans.append(_span(content, bold=bold, italic=italic,
                               underline=underline, subscript=True))
            continue

        # Normal character
        buf += c
        i += 1

    flush()


def _resolve_math(text: str) -> str:
    """Replace LaTeX math command tokens with Unicode, longest-match first."""
    # Handle sequences like \sharp\flat
    for cmd, sym in sorted(MATH_SYMBOL_MAP.items(), key=lambda x: -len(x[0])):
        text = text.replace(cmd, sym)
    return text


def parse_cell_content(raw: str) -> list:
    """
    Parse a raw LaTeX cell string into a list of span dicts.

    Each span dict has keys:
        text, bold, italic, underline, superscript, subscript

    Example
    -------
    >>> parse_cell_content(r"WebShaper-32B$^{\\dagger}$ \\textsuperscript{[17]}")
    [
      {'text': 'WebShaper-32B', 'bold': False, ...},
      {'text': '†',            'bold': False, 'superscript': True, ...},
      {'text': ' [17]',        'bold': False, 'superscript': True, ...},
    ]
    """
    spans: list = []
    _parse_segment(raw.strip(), spans)
    return _merge_spans(spans)


# ---------------------------------------------------------------------------
# LaTeX table structure parser
# ---------------------------------------------------------------------------

def _parse_tabular_spec(spec: str) -> list:
    """
    Extract column alignments from a tabular spec string.

    Handles:
      l / c / r            basic alignments
      p{} / m{} / b{}      paragraph columns  → 'l'
      X / L                tabularx/tabulary wide columns  → 'l'
      C                    tabulary centred  → 'c'
      R                    tabulary right  → 'r'
      | / || / :           vertical rules  (skipped)
      @{} / !{} / >{} / <{}  column decorators  (skipped, brace arg consumed)
      *{n}{sub_spec}       repeated spec  (recursively expanded)
    """
    aligns = []
    i = 0
    n = len(spec)

    while i < n:
        c = spec[i]

        # --- skip whitespace ---
        if c in " \t\n":
            i += 1
            continue

        # --- vertical rules and decorators without args ---
        if c in "|:":
            i += 1
            continue

        # --- basic alignments ---
        if c == 'l':
            aligns.append('l'); i += 1; continue
        if c == 'c':
            aligns.append('c'); i += 1; continue
        if c == 'r':
            aligns.append('r'); i += 1; continue

        # --- tabularx / tabulary column types ---
        if c == 'X' or c == 'L':
            aligns.append('l'); i += 1; continue
        if c == 'C':
            aligns.append('c'); i += 1; continue
        if c == 'R':
            aligns.append('r'); i += 1; continue

        # --- paragraph columns  p{} / m{} / b{} ---
        if c in ('p', 'm', 'b') and i + 1 < n and spec[i + 1] == '{':
            aligns.append('l')
            i += 1                          # skip 'p'/'m'/'b'
            i = _spec_skip_braced(spec, i)  # skip {width}
            continue

        # --- decorators that consume one braced arg: @{} !{} >{} <{} ---
        if c in ('@', '!', '>', '<') and i + 1 < n and spec[i + 1] == '{':
            i += 1
            i = _spec_skip_braced(spec, i)
            continue

        # --- *{n}{sub_spec} repeated column spec ---
        if c == '*':
            i += 1
            # Read {n}
            if i < n and spec[i] == '{':
                count_str, i = _spec_read_braced(spec, i)
                try:
                    count = int(count_str.strip())
                except ValueError:
                    count = 1
            else:
                count = 1
            # Read {sub_spec}
            if i < n and spec[i] == '{':
                sub_spec, i = _spec_read_braced(spec, i)
            else:
                sub_spec = ""
            sub_aligns = _parse_tabular_spec(sub_spec)
            aligns.extend(sub_aligns * count)
            continue

        # --- anything else: skip ---
        i += 1

    return aligns


def _spec_skip_braced(spec: str, i: int) -> int:
    """Advance past a brace-balanced {...} starting at spec[i] == '{'."""
    if i >= len(spec) or spec[i] != '{':
        return i
    depth = 0
    while i < len(spec):
        if spec[i] == '{':
            depth += 1
        elif spec[i] == '}':
            depth -= 1
            if depth == 0:
                return i + 1
        i += 1
    return i


def _spec_read_braced(spec: str, i: int):
    """
    Read the content of {...} at spec[i] == '{'.
    Returns (content: str, new_index: int).
    """
    if i >= len(spec) or spec[i] != '{':
        return "", i
    depth = 0
    start = i + 1
    while i < len(spec):
        if spec[i] == '{':
            depth += 1
        elif spec[i] == '}':
            depth -= 1
            if depth == 0:
                return spec[start:i], i + 1
        i += 1
    return spec[start:], len(spec)


def _hex_from_rgb_str(rgb_str: str) -> str:
    """Convert '245, 245, 245' → 'F5F5F5'."""
    try:
        parts = [int(x.strip()) for x in rgb_str.split(",")]
        return "{:02X}{:02X}{:02X}".format(*parts[:3])
    except ValueError:
        return "FFFFFF"


def _extract_cellcolor(text: str):
    """
    Look for \\cellcolor[RGB]{r,g,b} or \\cellcolor[HTML]{RRGGBB} at the start of text.
    Returns (hex_color: str | None, remaining_text: str).
    """
    m = re.match(
        r"\s*\\(?:cellcolor|rowcolor)\[RGB\]\{([^}]+)\}(.*)",
        text, re.DOTALL
    )
    if m:
        return _hex_from_rgb_str(m.group(1)), m.group(2).strip()

    m = re.match(
        r"\s*\\(?:cellcolor|rowcolor)\[HTML\]\{([0-9A-Fa-f]{6})\}(.*)",
        text, re.DOTALL
    )
    if m:
        return m.group(1).upper(), m.group(2).strip()

    m = re.match(
        r"\s*\\(?:cellcolor|rowcolor)\{([^}]+)\}(.*)",
        text, re.DOTALL
    )
    if m:
        # Named color — skip, return no color
        return None, m.group(2).strip()

    return None, text


def _split_cells(row_text: str) -> list:
    """
    Split a row string by '&' while respecting nested braces.
    Returns a list of raw cell strings.
    """
    cells = []
    depth = 0
    buf = ""
    i = 0
    while i < len(row_text):
        c = row_text[i]
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
        elif c == "&" and depth == 0:
            cells.append(buf)
            buf = ""
            i += 1
            continue
        buf += c
        i += 1
    cells.append(buf)
    return cells


def _classify_row(block: str):
    """
    Decide what kind of row a block of LaTeX text represents.

    Returns one of:
        ('skip', None)          — structural markers like \\toprule
        ('group', dict)         — \\multicolumn group header
        ('data', list)          — regular data row
    """
    b = block.strip()

    # Pure structural markers
    if re.match(r"^\\(toprule|bottomrule|midrule|hline)(\s*\\(toprule|bottomrule|midrule|hline))*\s*$", b):
        return "skip", None

    # Multicolumn group header
    mc = re.match(r"\\multicolumn\{(\d+)\}\{[^}]*\}\{(.*)\}\s*$", b, re.DOTALL)
    if mc:
        span = int(mc.group(1))
        inner = mc.group(2).strip()
        # May be wrapped in another brace layer from the regex
        color, text = _extract_cellcolor(inner)
        if color is None:
            color = "F5F5F5"
        # Strip any trailing braces left by imperfect brace matching
        text = re.sub(r"^\{|\}$", "", text).strip()
        return "group", {
            "text":     text,
            "bg_color": color,
            "span":     span,
        }

    # Rows that contain & (data rows)
    # Strip leading midrule/hline if any
    b_clean = re.sub(r"\\(midrule|hline)\s*", "", b).strip()
    if "&" in b_clean or b_clean:
        # Collapse multi-line row into a single line (some rows wrap onto next line)
        b_clean = re.sub(r"\s*\n\s*", " ", b_clean).strip()
        # Remove trailing \\ (LaTeX row break)
        b_clean = re.sub(r"\\\\$", "", b_clean).strip()
        cells_raw = _split_cells(b_clean)
        return "data", [c.strip() for c in cells_raw]

    return "skip", None


def _is_structural_only(line: str) -> bool:
    return bool(re.match(
        r"^\s*\\(toprule|bottomrule|midrule|hline|cline\{[^}]+\})\s*$",
        line
    ))



def _extract_caption_text(latex_str: str) -> str:
    """
    Extract the content of \\caption{...} using brace-depth tracking so that
    captions containing nested braces (e.g. \\textbf{foo}) are captured in full,
    regardless of whether the caption appears before or after the tabular body.

    Returns an empty string when no caption is found.
    """
    for m in re.finditer(r"\\caption\*?\s*", latex_str):
        i = m.end()
        n = len(latex_str)
        # Skip optional [short caption]
        while i < n and latex_str[i] in " \t\n":
            i += 1
        if i < n and latex_str[i] == "[":
            while i < n and latex_str[i] != "]":
                i += 1
            i += 1
        while i < n and latex_str[i] in " \t\n":
            i += 1
        if i >= n or latex_str[i] != "{":
            continue
        content, _ = _spec_read_braced(latex_str, i)
        return content.strip()
    return ""


def _extract_tabular_spec_and_body(latex_str: str):
    """
    Locate \\begin{tabular}…\\end{tabular} (or tabular*) in *latex_str* and
    return (spec_string, body_string).

    Uses a brace-depth counter so specs containing nested braces like
    p{0.38\\textwidth} are handled correctly.  Returns (None, None) when
    no tabular environment is found.
    """
    # Find the start of \begin{tabular} or \begin{tabular*}
    begin_m = re.search(r"\\begin\{tabular\*?\}", latex_str)
    if not begin_m:
        return None, None

    i = begin_m.end()
    n = len(latex_str)

    # Skip optional [pos] argument like \begin{tabular}[t]{spec}
    while i < n and latex_str[i] in " \t\n":
        i += 1
    if i < n and latex_str[i] == "[":
        while i < n and latex_str[i] != "]":
            i += 1
        i += 1  # skip ']'

    # Read the {spec} argument with full brace-depth tracking
    while i < n and latex_str[i] in " \t\n":
        i += 1
    if i >= n or latex_str[i] != "{":
        return None, None

    spec, i = _spec_read_braced(latex_str, i)

    # The body runs until the matching \end{tabular} / \end{tabular*}
    # (use the same env name that was opened)
    env_name = re.search(r"\\begin\{(tabular\*?)\}", latex_str).group(1)
    end_tag = f"\\end{{{env_name}}}"
    end_pos = latex_str.find(end_tag, i)
    if end_pos == -1:
        body = latex_str[i:]
    else:
        body = latex_str[i:end_pos]

    return spec, body


def parse_latex_table(latex_str: str) -> dict:
    """
    Parse a LaTeX table block and return a structured representation.

    Accepts both the full \\begin{table}...\\end{table} wrapper and a bare
    \\begin{tabular}...\\end{tabular} block.

    Returns
    -------
    {
        'caption':   str,
        'col_count': int,
        'col_align': ['l'|'c'|'r', ...],
        'rows': [
            {'type': 'header', 'cells': [[span, ...], ...]},
            {'type': 'group',  'text': str, 'bg_color': str, 'span': int},
            {'type': 'data',   'cells': [[span, ...], ...]},
        ]
    }
    """
    result = {
        "caption":   "",
        "col_count": 0,
        "col_align": [],
        "rows":      [],
    }

    # --- Caption (brace-depth aware to handle nested braces) ---
    result["caption"] = _extract_caption_text(latex_str)

    # --- Tabular spec + body (brace-depth aware) ---
    # The simple [^}]+ regex breaks on specs like p{0.38\textwidth}.
    # We scan manually using a brace-depth counter so nested braces are
    # handled correctly.
    spec_raw, body = _extract_tabular_spec_and_body(latex_str)
    if spec_raw is None:
        return result

    aligns = _parse_tabular_spec(spec_raw)
    result["col_align"] = aligns
    result["col_count"] = len(aligns)

    if body is None:
        return result

    # Normalise line endings.
    # Convert LaTeX row-break \\\\ into a blank-line separator so every
    # \\-terminated row becomes its own block for the blank-line splitter.
    body = body.replace("\r\n", "\n").replace("\r", "\n")
    body = re.sub(r"\\\\\s*\n", "\n\n", body)  # \\ row-break → blank line

    # Split body into "blocks" separated by blank lines.
    # A block corresponds to one logical row (which may span several lines
    # if the cell content is long).
    raw_blocks = re.split(r"\n[ \t]*\n", body.strip())

    # Second pass: blocks may contain multiple sub-rows if they contain \midrule
    # in the middle.  Split those up.
    blocks = []
    for rb in raw_blocks:
        rb = rb.strip()
        if not rb:
            continue
        # Split at mid-row \midrule lines
        sub = re.split(r"\n[ \t]*\\(midrule|hline)[ \t]*\n", rb)
        for s in sub:
            s = s.strip()
            if s and s not in ("midrule", "hline"):
                blocks.append(s)

    first_data_seen = False

    for block in blocks:
        # Skip pure structural-only blocks
        if all(_is_structural_only(ln) for ln in block.splitlines() if ln.strip()):
            continue

        # Remove leading structural markers within the block
        clean = "\n".join(
            ln for ln in block.splitlines()
            if not _is_structural_only(ln)
        ).strip()
        if not clean:
            continue

        kind, payload = _classify_row(clean)

        if kind == "skip":
            continue

        if kind == "group":
            result["rows"].append({"type": "group", **payload})
            first_data_seen = True
            continue

        if kind == "data":
            cells_raw = payload
            n = result["col_count"] or len(cells_raw)
            # Pad or truncate to column count
            while len(cells_raw) < n:
                cells_raw.append("")

            parsed_cells = [parse_cell_content(c) for c in cells_raw[:n]]

            row_type = "header" if not first_data_seen else "data"
            first_data_seen = True
            result["rows"].append({"type": row_type, "cells": parsed_cells})

    return result


# ---------------------------------------------------------------------------
# python-docx rendering helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, fill_hex: str):
    """Apply a solid background color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.upper())
    # Remove existing shd if any to avoid duplicates
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    """Set inner padding for a table cell (twips)."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table):
    """Apply single-line borders to all sides and inner dividers."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def _set_col_widths(table, widths_cm):
    """Set column widths in centimetres."""
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


def _apply_font(run, size_pt=10, font_name="Times New Roman",
                east_asia_font="宋体"):
    """Apply standard font settings to a run."""
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    r_el = run._element
    if r_el.rPr is None:
        r_el.get_or_add_rPr()
    r_el.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def _fill_cell_from_spans(cell, spans: list, align: str = "center",
                          font_pt: int = 10):
    """
    Write a list of span dicts into a table cell, applying all formatting.

    Parameters
    ----------
    cell     : docx table cell
    spans    : list of span dicts produced by parse_cell_content()
    align    : 'left' | 'center' | 'right'
    font_pt  : base font size in points
    """
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = {
        "left":  WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    for s in spans:
        run = para.add_run(s["text"])
        run.bold      = s["bold"]
        run.italic    = s["italic"]
        run.underline = s["underline"]
        run.font.superscript = s["superscript"]
        run.font.subscript   = s["subscript"]
        # Superscripts use a slightly smaller size
        pt = font_pt - 2 if (s["superscript"] or s["subscript"]) else font_pt
        _apply_font(run, size_pt=pt)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)


def _col_widths_cm(col_count: int, first_col_cm: float = None,
                   other_col_cm: float = 2.0,
                   total_cm: float = 16.8) -> list:
    """
    Compute a sensible list of column widths.

    If *first_col_cm* is None, the first column gets the remaining space after
    the other columns are each allocated *other_col_cm*.
    """
    if col_count <= 0:
        return []
    if col_count == 1:
        return [total_cm]
    if first_col_cm is None:
        first_col_cm = max(2.0, total_cm - (col_count - 1) * other_col_cm)
    return [first_col_cm] + [other_col_cm] * (col_count - 1)


# ---------------------------------------------------------------------------
# Caption renderer
# ---------------------------------------------------------------------------

def _render_caption(doc, caption_latex: str):
    """Add a caption paragraph above the table (楷体, centred, italic)."""
    if not caption_latex.strip():
        return
    spans = parse_cell_content(caption_latex)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for s in spans:
        run = p.add_run(s["text"])
        run.italic    = True          # captions are always italic
        run.bold      = s["bold"]
        run.underline = s["underline"]
        run.font.superscript = s["superscript"]
        pt = 8 if s["superscript"] else 10
        _apply_font(run, size_pt=pt,
                    font_name="Times New Roman",
                    east_asia_font="楷体")


# ---------------------------------------------------------------------------
# Main renderer
# ---------------------------------------------------------------------------

def render_latex_table_to_docx(doc, latex_str: str,
                                col_widths: list = None,
                                font_pt: int = 10) -> None:
    """
    Parse *latex_str* and append a formatted table (and optional caption) to *doc*.

    Parameters
    ----------
    doc         : docx.Document
    latex_str   : LaTeX source containing \\begin{table} or \\begin{tabular}
    col_widths  : optional list of column widths in cm; auto-computed if None
    font_pt     : base font size in points (default 10)
    """
    data = parse_latex_table(latex_str)

    col_count = data["col_count"]
    if col_count == 0:
        # Fallback: try to infer from the widest row
        for row in data["rows"]:
            if row["type"] in ("header", "data"):
                col_count = len(row["cells"])
                break
    if col_count == 0:
        return

    col_aligns = data["col_align"] or ["l"] + ["c"] * (col_count - 1)

    # --- Caption ---
    # _render_caption(doc, data["caption"])

    # --- Column widths ---
    if col_widths is None:
        col_widths = _col_widths_cm(col_count)

    # --- Create table ---
    table = doc.add_table(rows=0, cols=col_count)
    table.alignment   = WD_TABLE_ALIGNMENT.CENTER
    table.style       = "Table Grid"
    table.autofit     = False
    _set_table_borders(table)

    for row_data in data["rows"]:
        rtype = row_data["type"]

        # ---- Header row ----
        if rtype == "header":
            row = table.add_row()
            for ci, cell_spans in enumerate(row_data["cells"]):
                align = col_aligns[ci] if ci < len(col_aligns) else "c"
                align_str = {"l": "left", "r": "right"}.get(align, "center")
                # Header cells are always centered and bold
                bold_spans = [dict(s, bold=True) for s in cell_spans]
                _fill_cell_from_spans(row.cells[ci], bold_spans,
                                      align="center", font_pt=font_pt)

        # ---- Group header (multicolumn with background) ----
        elif rtype == "group":
            row = table.add_row()
            span = min(row_data["span"], col_count)
            # Merge cells across the row
            merged = row.cells[0]
            for ci in range(1, span):
                merged = merged.merge(row.cells[ci])
            # Set background color
            _set_cell_shading(row.cells[0], row_data["bg_color"])
            # Render text as bold centered
            spans = parse_cell_content(row_data["text"])
            bold_spans = [dict(s, bold=True) for s in spans]
            _fill_cell_from_spans(row.cells[0], bold_spans,
                                  align="center", font_pt=font_pt)

        # ---- Regular data row ----
        elif rtype == "data":
            row = table.add_row()
            for ci, cell_spans in enumerate(row_data["cells"]):
                if ci >= col_count:
                    break
                align = col_aligns[ci] if ci < len(col_aligns) else "c"
                align_str = {"l": "left", "r": "right"}.get(align, "center")
                _fill_cell_from_spans(row.cells[ci], cell_spans,
                                      align=align_str, font_pt=font_pt)

    # Apply column widths (must be done after all rows are added)
    _set_col_widths(table, col_widths)

    # Ensure vertical centering on every cell
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
