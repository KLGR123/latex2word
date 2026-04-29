from __future__ import annotations

import copy
import os
import subprocess
import tempfile
import zipfile
from typing import Optional, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

from .context import RenderContext
from .footnotes import merge_pandoc_footnotes
from .settings import TEX_POSTAMBLE, TEX_PREAMBLE, NS_W
from .word_xml import clean_xml_text_nodes, fix_lxml_run_fonts, fix_math_text_runs, indent, set_para_font, spacing


def pandoc_to_xml(latex_body: str) -> Tuple[Optional[bytes], Optional[bytes]]:
    src = TEX_PREAMBLE + latex_body + TEX_POSTAMBLE
    tex_path = None
    docx_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".tex", mode="w", encoding="utf-8", delete=False) as handle:
            handle.write(src)
            tex_path = handle.name
        docx_path = tex_path.replace(".tex", ".docx")
        result = subprocess.run(
            ["pandoc", tex_path, "-f", "latex", "-t", "docx", "-o", docx_path, "--quiet"],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode != 0:
            print(f"  [pandoc] {result.stderr[:200]}")
            return None, None
        with zipfile.ZipFile(docx_path) as archive:
            names = archive.namelist()
            return (
                archive.read("word/document.xml"),
                archive.read("word/footnotes.xml") if "word/footnotes.xml" in names else None,
            )
    except Exception as exc:
        print(f"  [pandoc] {exc}")
        return None, None
    finally:
        for path in (tex_path, docx_path):
            if path:
                try:
                    os.unlink(path)
                except OSError:
                    pass


def inject_pandoc_paras(
    latex_body: str,
    doc,
    ctx: RenderContext,
    font_name: str,
    size_pt: float,
    align: WD_ALIGN_PARAGRAPH,
    indent_first_cm: float = 0.0,
    indent_left_cm: float = 0.0,
    before_pt: float = 0.0,
    after_pt: float = 6.0,
) -> bool:
    doc_xml, fn_xml = pandoc_to_xml(latex_body)
    if doc_xml is None:
        return False

    tree = etree.fromstring(doc_xml)
    merge_pandoc_footnotes(tree, fn_xml, ctx)

    injected = 0
    for src_p in tree.findall(f".//{{{NS_W}}}p"):
        if all(etree.QName(child.tag).localname == "pPr" for child in src_p):
            continue
        dst = doc.add_paragraph()
        dst.alignment = align
        indent(dst, left_cm=indent_left_cm, first_line_cm=indent_first_cm)
        spacing(dst, before_pt=before_pt, after_pt=after_pt)
        set_para_font(dst, font_name, size_pt)
        for child in src_p:
            if etree.QName(child.tag).localname == "pPr":
                continue
            node = copy.deepcopy(child)
            if etree.QName(node.tag).localname == "r":
                fix_lxml_run_fonts(node, font_name, size_pt)
            fix_math_text_runs(node, font_name)
            dst._p.append(node)
        injected += 1

    clean_xml_text_nodes(doc.paragraphs[len(doc.paragraphs) - injected:])
    return injected > 0
