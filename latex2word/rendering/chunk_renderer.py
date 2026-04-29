from __future__ import annotations

import re
import sys
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .classifier import classify_paragraph
from .context import RenderContext
from .latex_utils import (
    clean_cross_refs,
    clean_spaces,
    clean_title_latex,
    extract_caption,
    extract_display_math_block,
    extract_env_body,
    has_wrapped_display_math,
    normalize_math_fonts,
    sanitize_fragment_for_pandoc,
    split_environment_fragment,
    strip_equation_trailing_punct,
)
from .pandoc_bridge import inject_pandoc_paras
from .settings import COLORS, FONTS, MATH_ENV_LABELS, RE_MATHENV, SIZES
from .table import parse_latex_table, render_latex_table_to_docx
from .word_xml import indent, set_para_font, set_run_font, shade, spacing

try:
    import cn2an as _cn2an

    def to_chinese_number(n: int) -> str:
        return _cn2an.an2cn(str(n))
except ImportError:
    def to_chinese_number(n: int) -> str:
        return str(n)


class ChunkRenderer:
    def __init__(self, ctx: RenderContext):
        self.ctx = ctx

    def render_para(
        self,
        doc,
        latex_text: str,
        font_name: str = FONTS.body,
        size_pt: float = SIZES.body,
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
        indent_first_cm: float = 0.74,
        indent_left_cm: float = 0.0,
        before_pt: float = 0.0,
        after_pt: float = 6.0,
    ) -> bool:
        text = sanitize_fragment_for_pandoc(
            normalize_math_fonts(clean_cross_refs(latex_text)).strip()
        )
        if not text:
            return True
        ok = inject_pandoc_paras(
            text,
            doc,
            self.ctx,
            font_name=font_name,
            size_pt=size_pt,
            align=align,
            indent_first_cm=indent_first_cm,
            indent_left_cm=indent_left_cm,
            before_pt=before_pt,
            after_pt=after_pt,
        )
        if ok:
            return True
        plain = re.sub(r'\\[a-zA-Z]+\*?\{([^{}]*)\}', r'\1', text)
        plain = re.sub(r'\\[a-zA-Z]+\*?', '', plain)
        plain = plain.replace('$', '').replace('{', '').replace('}', '').strip()
        p = doc.add_paragraph()
        p.alignment = align
        indent(p, left_cm=indent_left_cm, first_line_cm=indent_first_cm)
        spacing(p, before_pt=before_pt, after_pt=after_pt)
        set_run_font(p.add_run(plain), font_name, size_pt)
        return False

    def render_chapter_title(self, doc, meta: dict) -> None:
        chapter = meta.get("chapter", "")
        title_zh = meta.get("title_translation", meta.get("title", ""))
        title_zh = re.sub(r'^#+\s+', '', title_zh)
        title_zh = clean_title_latex(title_zh)
        try:
            chapter_cn = to_chinese_number(int(chapter))
        except (ValueError, TypeError):
            chapter_cn = str(chapter)
        heading = doc.add_heading(f"第{chapter_cn}章  {title_zh}", level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            set_run_font(run, FONTS.heading, SIZES.h1 + 2, bold=True)
        spacing(heading, before_pt=24, after_pt=12)

    @staticmethod
    def extract_heading_text(translation: str, cmd: str) -> Tuple[str, str]:
        match = re.match(r'\\' + re.escape(cmd) + r'\*?\{', translation.strip())
        if not match:
            return "", translation
        depth = 0
        idx = match.end() - 1
        while idx < len(translation):
            if translation[idx] == "{":
                depth += 1
            elif translation[idx] == "}":
                depth -= 1
                if depth == 0:
                    return translation[match.end():idx].strip(), translation[idx + 1:].strip()
            idx += 1
        return translation[match.end():].strip(), ""

    def render_inline_heading(self, doc, para: dict) -> None:
        translation = para.get("translation", para.get("text", "")).strip()
        title, body = self.extract_heading_text(translation, "paragraph")
        if not title:
            title, body = self.extract_heading_text(translation, "subparagraph")
        if not title:
            self.render_paragraph(doc, para)
            return
        title_clean = re.sub(r'\\[a-zA-Z]+\*?\{([^{}]*)\}', r'\1', title)
        title_clean = re.sub(r'\\[a-zA-Z]+\*?', '', title_clean).strip()
        body_clean = re.sub(r'^[\s。，、；：！？…—·]+', '', body.strip())
        body_clean = clean_spaces(body_clean)
        combined = f"\\textbf{{{title_clean}}} {body_clean}" if body_clean else f"\\textbf{{{title_clean}}}"
        before = len(doc.paragraphs)
        ok = inject_pandoc_paras(
            combined,
            doc,
            self.ctx,
            font_name=FONTS.body,
            size_pt=SIZES.body,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            before_pt=6.0,
            after_pt=4.0,
        )
        for p in doc.paragraphs[before:]:
            indent(p, first_line_cm=0.74)
        if not ok:
            p = doc.add_paragraph()
            spacing(p, before_pt=6, after_pt=4)
            indent(p, first_line_cm=0.74)
            set_run_font(p.add_run(title_clean), FONTS.body, SIZES.body, bold=True)
            if body_clean:
                set_run_font(p.add_run(body_clean), FONTS.body, SIZES.body)

    def render_heading(self, doc, para: dict, level: int) -> None:
        translation = para.get("translation", para.get("text", "")).strip()
        env_label = para.get("env_label", "")
        cmd_map = {1: "section", 2: "subsection", 3: "subsubsection"}
        heading_text, body_text = self.extract_heading_text(translation, cmd_map[level])
        if not heading_text:
            match = re.search(r'\{([^}]+)\}', translation)
            heading_text = match.group(1) if match else translation
            body_text = ""
        heading_clean = re.sub(r'\\[a-zA-Z]+\*?\{([^{}]*)\}', r'\1', heading_text)
        heading_clean = re.sub(r'\\[a-zA-Z]+\*?', '', heading_clean).strip()
        heading_clean = clean_spaces(heading_clean)
        heading_clean = re.sub(r'^#+\s+', '', heading_clean)
        numeric = re.sub(r'[节小].*$', '', env_label)
        prefix = (numeric + "  ") if re.match(r'^\d', numeric) else ""
        heading = doc.add_heading(prefix + heading_clean, level=level)
        for run in heading.runs:
            set_run_font(run, FONTS.heading, {1: SIZES.h1, 2: SIZES.h2, 3: SIZES.h3}[level], bold=True)
        spacing(heading, before_pt=14, after_pt=5)
        if body_text.strip():
            self.render_para(doc, clean_spaces(body_text))

    def render_abstract(self, doc, para: dict) -> None:
        translation = para.get("translation", para.get("text", ""))
        _, body = extract_env_body(translation, "abstract")
        if not body:
            body = translation
        heading = doc.add_heading("摘要", level=1)
        for run in heading.runs:
            set_run_font(run, FONTS.heading, SIZES.h1, bold=True)
        spacing(heading, before_pt=14, after_pt=6)
        self.render_para(doc, clean_spaces(body), indent_first_cm=0.74, indent_left_cm=1.0)

    def render_paragraph(self, doc, para: dict) -> None:
        self.render_para(doc, clean_spaces(para.get("translation", para.get("text", ""))))

    def render_equation(self, doc, para: dict) -> None:
        translation = para.get("translation", para.get("text", "")).strip()
        env_label = para.get("env_label", "")
        equation_source = extract_display_math_block(translation) if has_wrapped_display_math(translation) else translation
        ok = inject_pandoc_paras(
            strip_equation_trailing_punct(clean_cross_refs(equation_source)),
            doc,
            self.ctx,
            font_name=FONTS.body,
            size_pt=SIZES.body,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            before_pt=6.0,
            after_pt=2.0,
        )
        if not ok:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacing(p, before_pt=6, after_pt=2)
            run = p.add_run(equation_source)
            run.font.name = FONTS.code
            run.font.size = Pt(SIZES.code)
        if env_label:
            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            spacing(lp, before_pt=0, after_pt=8)
            set_run_font(lp.add_run(f"（{env_label}）"), FONTS.body, 10)

    def add_caption(self, doc, label: str, caption_latex: str) -> None:
        if not label and not caption_latex:
            return
        cleaned = clean_spaces(caption_latex) if caption_latex else ""
        full_latex = f"\\textbf{{{label}}} {cleaned}" if cleaned else f"\\textbf{{{label}}}"
        ok = inject_pandoc_paras(
            full_latex,
            doc,
            self.ctx,
            font_name=FONTS.caption,
            size_pt=SIZES.caption,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            before_pt=2.0,
            after_pt=8.0,
        )
        if not ok:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacing(p, before_pt=2, after_pt=8)
            parts = re.split(r'(\$[^$]+\$)', f"{label} {cleaned}")
            for part in parts:
                if part.startswith("$") and part.endswith("$"):
                    math_ok = inject_pandoc_paras(
                        part,
                        doc,
                        self.ctx,
                        font_name=FONTS.caption,
                        size_pt=SIZES.caption,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                    )
                    if not math_ok:
                        set_run_font(p.add_run(part), FONTS.caption, SIZES.caption)
                else:
                    set_run_font(p.add_run(part), FONTS.caption, SIZES.caption, italic=False)

    @staticmethod
    def find_minipages(text: str) -> List[Dict[str, Any]]:
        results = []
        pattern = re.compile(
            r'\\begin\{(?:minipage|subfigure|subfloat)\}'
            r'(?:\[[^\]]*\])?(?:\{[^}]*\})?'
            r'(.*?)'
            r'\\end\{(?:minipage|subfigure|subfloat)\}',
            re.DOTALL,
        )
        for match in pattern.finditer(text):
            content = match.group(1)
            images = re.findall(r'\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}', content)
            caption = extract_caption(content)
            if images or caption:
                results.append({"images": images, "caption": caption})
        return results

    @staticmethod
    def remove_table_borders(table) -> None:
        tbl_pr = table._tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr)
        old = tbl_pr.find(qn("w:tblBorders"))
        if old is not None:
            tbl_pr.remove(old)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "none")
            borders.append(el)
        tbl_pr.append(borders)

    def pdf_to_image(self, pdf_path: Path) -> Optional[Path]:
        print(f"  [debug] python = {sys.executable}")
        try:
            import fitz  # type: ignore

            pdf_doc = fitz.open(str(pdf_path))
            page = pdf_doc[0]
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat)
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            pix.save(tmp.name)
            tmp.close()
            pdf_doc.close()
            self.ctx.temp_files.append(tmp.name)
            print(f"  [pdf→png] {pdf_path.name} -> {tmp.name} (fitz)")
            return Path(tmp.name)
        except ImportError:
            pass
        except Exception as exc:
            import traceback

            print(f"  [warn] fitz PDF conversion failed for {pdf_path}: {exc}")
            traceback.print_exc()

        try:
            from pdf2image import convert_from_path  # type: ignore

            images = convert_from_path(str(pdf_path), first_page=1, last_page=1, dpi=144)
            if images:
                tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                images[0].save(tmp.name, "PNG")
                tmp.close()
                self.ctx.temp_files.append(tmp.name)
                print(f"  [pdf→png] {pdf_path.name} -> {tmp.name} (pdf2image)")
                return Path(tmp.name)
        except ImportError:
            pass
        except Exception as exc:
            print(f"  [warn] pdf2image conversion failed for {pdf_path}: {exc}")

        print(f"  [warn] No PDF-to-image library found; cannot embed {pdf_path.name}.")
        print("         Install PyMuPDF:  pip install pymupdf")
        print("         or pdf2image:     pip install pdf2image")
        return None

    def resolve_image(self, figures_dir: str, rel: str) -> Optional[Path]:
        base = Path(figures_dir) / rel
        for candidate in [base, *[base.with_suffix(s) for s in (".png", ".jpg", ".jpeg", ".pdf")]]:
            if candidate.exists():
                if candidate.suffix.lower() == ".pdf":
                    return self.pdf_to_image(candidate)
                return candidate
        return None

    def render_figure(self, doc, para: dict, figures_dir: str = ".") -> None:
        translation = para.get("translation", para.get("text", ""))
        env_label = para.get("env_label", "")
        if re.search(r'\\begin\{(?:minipage|subfigure|subfloat)\}', translation):
            minipages = self.find_minipages(translation)
            if minipages:
                count = len(minipages)
                table = doc.add_table(rows=2, cols=count)
                self.remove_table_borders(table)
                col_w = Inches(6.0 / count * 0.9)
                for col, mp in enumerate(minipages):
                    img_para = table.cell(0, col).paragraphs[0]
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    embedded = False
                    for rel in mp["images"]:
                        found = self.resolve_image(figures_dir, rel)
                        if found:
                            try:
                                img_para.add_run().add_picture(str(found), width=col_w)
                                embedded = True
                                break
                            except Exception as exc:
                                print(f"  [warn] subfigure {found}: {exc}")
                    if not embedded and mp["images"]:
                        set_run_font(
                            img_para.add_run(f"[图片: {mp['images'][0]}]"),
                            FONTS.body,
                            9,
                            color=RGBColor(0x88, 0x88, 0x88),
                        )
                    if mp["caption"]:
                        cap_para = table.cell(1, col).paragraphs[0]
                        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_run_font(cap_para.add_run(clean_spaces(mp["caption"])), FONTS.caption, SIZES.caption)
                lp = doc.add_paragraph()
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                spacing(lp, before_pt=4, after_pt=8)
                if env_label:
                    set_run_font(lp.add_run(env_label), FONTS.caption, SIZES.caption, bold=True)
                return

        image_paths = re.findall(r'\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}', translation)
        caption_latex = extract_caption(translation)
        embedded = False
        for rel in image_paths:
            found = self.resolve_image(figures_dir, rel)
            if found:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(found), width=Inches(5.0))
                    embedded = True
                except Exception as exc:
                    print(f"  [warn] figure {found}: {exc}")
        if not embedded and image_paths:
            ph = doc.add_paragraph()
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade(ph, "F2F2F2")
            set_run_font(
                ph.add_run(f"[图片: {', '.join(image_paths)}]"),
                FONTS.body,
                10,
                color=RGBColor(0x88, 0x88, 0x88),
            )
        self.add_caption(doc, env_label, caption_latex)

    @staticmethod
    def set_table_auto_col_widths(doc, total_cm: float = 15.5) -> None:
        if not doc.tables:
            return
        tbl = doc.tables[-1]
        n_cols = len(tbl.columns)
        if n_cols == 0:
            return
        col_chars = [0] * n_cols
        for row in tbl.rows:
            cells = row.cells
            for i in range(min(n_cols, len(cells))):
                length = sum(len(p.text) for p in cells[i].paragraphs)
                if length > col_chars[i]:
                    col_chars[i] = length
        total_chars = sum(col_chars)
        total_twips = int(total_cm * 567)
        min_twips = total_twips // (n_cols * 10)
        if total_chars == 0:
            col_twips = [total_twips // n_cols] * n_cols
        else:
            raw = [int(total_twips * c / total_chars) for c in col_chars]
            col_twips = [max(w, min_twips) for w in raw]
            excess = sum(col_twips) - total_twips
            if excess != 0:
                col_twips[col_twips.index(max(col_twips))] -= excess
        tbl_pr = tbl._tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl._tbl.insert(0, tbl_pr)
        for old_tag in (qn("w:tblW"), qn("w:tblLayout")):
            old = tbl_pr.find(old_tag)
            if old is not None:
                tbl_pr.remove(old)
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), str(total_twips))
        tbl_w.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_w)
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tbl_pr.append(layout)
        old_grid = tbl._tbl.find(qn("w:tblGrid"))
        if old_grid is not None:
            tbl._tbl.remove(old_grid)
        grid = OxmlElement("w:tblGrid")
        for width in col_twips:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(width))
            grid.append(gc)
        tbl_pr_idx = list(tbl._tbl).index(tbl_pr)
        tbl._tbl.insert(tbl_pr_idx + 1, grid)
        for row in tbl.rows:
            cells = row.cells
            for i in range(min(n_cols, len(cells))):
                tc = cells[i]._tc
                tc_pr = tc.find(qn("w:tcPr"))
                if tc_pr is None:
                    tc_pr = OxmlElement("w:tcPr")
                    tc.insert(0, tc_pr)
                old_tcw = tc_pr.find(qn("w:tcW"))
                if old_tcw is not None:
                    tc_pr.remove(old_tcw)
                tc_w = OxmlElement("w:tcW")
                tc_w.set(qn("w:w"), str(col_twips[i]))
                tc_w.set(qn("w:type"), "dxa")
                tc_pr.append(tc_w)

    def render_table(self, doc, para: dict) -> None:
        source = para.get("translation") or para.get("text", "")
        env_label = para.get("env_label", "")
        try:
            has_content = bool(parse_latex_table(source).get("rows"))
        except Exception:
            has_content = False
        if has_content:
            try:
                caption_latex = extract_caption(source)
                self.add_caption(doc, env_label, caption_latex)
                render_latex_table_to_docx(doc, source)
                self.set_table_auto_col_widths(doc)
                return
            except Exception as exc:
                print(f"  [table] Renderer failed for {env_label}: {exc}")
        header = doc.add_paragraph()
        shade(header, COLORS.table_bg)
        spacing(header, before_pt=6, after_pt=2)
        set_run_font(
            header.add_run(f"[表格 {env_label}]"),
            FONTS.body,
            10,
            bold=True,
            color=RGBColor(0x80, 0x60, 0x00),
        )
        snippet = source[:600] + ("…" if len(source) > 600 else "")
        for line in snippet.split("\n"):
            p = doc.add_paragraph()
            indent(p, left_cm=0.5)
            shade(p, COLORS.table_bg)
            spacing(p, before_pt=0, after_pt=0)
            run = p.add_run(line)
            run.font.name = FONTS.code
            run.font.size = Pt(8)
        caption_latex = extract_caption(source)
        if caption_latex:
            self.add_caption(doc, env_label, caption_latex)

    def render_code(self, doc, para: dict) -> None:
        text = para.get("text", "")
        env_label = para.get("env_label", "")
        body = text.strip()
        for env in ("verbatim", "lstlisting", "minted", "Verbatim"):
            _, inner = extract_env_body(body, env)
            if inner:
                body = inner
                break
        else:
            body = re.sub(r'^\\begin\{[^}]+\}\*?', '', body)
            body = re.sub(r'\\end\{[^}]+\}\*?$', '', body).strip()
        if env_label:
            lp = doc.add_paragraph()
            spacing(lp, before_pt=8, after_pt=2)
            set_run_font(lp.add_run(env_label), FONTS.body, SIZES.code, bold=True)
        for line in body.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            indent(p, left_cm=0.5)
            shade(p, COLORS.code_bg)
            run = p.add_run(line)
            run.font.name = FONTS.code
            run.font.size = Pt(SIZES.code)

    def render_mathenv(self, doc, para: dict) -> None:
        translation = para.get("translation", para.get("text", "")).strip()
        match = RE_MATHENV.match(translation)
        if not match:
            self.render_paragraph(doc, para)
            return
        env_name = match.group(1).lower().rstrip("*")
        label_zh = MATH_ENV_LABELS.get(env_name, env_name)
        opt, body = extract_env_body(translation, env_name)
        if not body:
            self.render_paragraph(doc, para)
            return
        display_label = f"【{label_zh}】" + (f"（{opt}）" if opt else "")
        header = doc.add_paragraph()
        shade(header, COLORS.theorem_bg)
        indent(header, left_cm=0.8)
        spacing(header, before_pt=8, after_pt=0)
        set_run_font(header.add_run(display_label), FONTS.body, SIZES.body, bold=True)
        ok = inject_pandoc_paras(
            clean_spaces(body),
            doc,
            self.ctx,
            font_name=FONTS.body,
            size_pt=SIZES.body,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            indent_left_cm=0.8,
            before_pt=2.0,
            after_pt=8.0,
        )
        if not ok:
            fp = doc.add_paragraph()
            shade(fp, COLORS.theorem_bg)
            indent(fp, left_cm=0.8)
            spacing(fp, before_pt=2, after_pt=8)
            set_run_font(fp.add_run(clean_spaces(body)), FONTS.body, SIZES.body, italic=True)

    def render_proof(self, doc, para: dict) -> None:
        translation = para.get("translation", para.get("text", "")).strip()
        before, body, after = split_environment_fragment(translation, "proof")
        opt, extracted_body = extract_env_body(translation, "proof")
        if extracted_body != translation:
            body = extracted_body
        label_zh = "证明" + (f"（{opt}）" if opt else "")

        if before.strip():
            self.render_para(
                doc,
                sanitize_fragment_for_pandoc(before, strip_proof_markers=True),
                before_pt=2.0,
                after_pt=4.0,
            )

        lp = doc.add_paragraph()
        indent(lp, left_cm=0.8)
        spacing(lp, before_pt=4, after_pt=0)
        set_run_font(lp.add_run(label_zh + "  "), FONTS.body, SIZES.body, bold=True)
        ok = inject_pandoc_paras(
            sanitize_fragment_for_pandoc(clean_spaces(body) if body else "", strip_proof_markers=True),
            doc,
            self.ctx,
            font_name=FONTS.body,
            size_pt=SIZES.body,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            indent_left_cm=0.8,
            before_pt=2.0,
            after_pt=4.0,
        )
        if not ok and body:
            fp = doc.add_paragraph()
            indent(fp, left_cm=0.8)
            spacing(fp, before_pt=2, after_pt=4)
            set_run_font(fp.add_run(clean_spaces(body)), FONTS.body, SIZES.body, italic=True)
        qp = doc.add_paragraph()
        qp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        spacing(qp, before_pt=0, after_pt=4)
        set_run_font(qp.add_run("□"), FONTS.body, SIZES.body)

        if after.strip():
            self.render_para(
                doc,
                sanitize_fragment_for_pandoc(after, strip_proof_markers=True),
                before_pt=2.0,
                after_pt=6.0,
            )

    def render_references(self, doc, citations_path: str) -> None:
        path = Path(citations_path)
        if not path.exists():
            print(f"  [warn] citations file not found: {citations_path}")
            return
        try:
            import json

            with open(path, encoding="utf-8") as handle:
                raw: Dict[str, Any] = json.load(handle)
        except Exception as exc:
            print(f"  [warn] Could not load citations.json: {exc}")
            return
        id_to_citation: Dict[int, str] = {}
        for entry in raw.values():
            try:
                nid = int(entry["id"])
                if nid not in id_to_citation:
                    id_to_citation[nid] = str(entry["citation"])
            except (KeyError, TypeError, ValueError):
                pass
        if not id_to_citation:
            print("  [warn] citations.json has no valid entries -- skipping.")
            return
        doc.add_page_break()
        heading = doc.add_heading("参考文献", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            set_run_font(run, FONTS.heading, SIZES.h1, bold=True)
        spacing(heading, before_pt=0, after_pt=10)
        size = SIZES.body - 1
        for nid in sorted(id_to_citation):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            indent(p, left_cm=1.2, hanging_cm=1.2)
            spacing(p, before_pt=0, after_pt=4)
            set_para_font(p, FONTS.body, size)
            set_run_font(p.add_run(f"[{nid}] "), FONTS.body, size, bold=True)
            set_run_font(p.add_run(clean_spaces(id_to_citation[nid])), FONTS.body, size)
        print(f"  [references] Appended {len(id_to_citation)} entries.")

    def render_chunk(self, doc, para: dict, figures_dir: str) -> None:
        if para.get("syntax_error", False):
            print(f"  [skip] id={para.get('id')} (syntax_error flag)")
            return
        chunk_type = classify_paragraph(para)
        print(f"  [{para.get('id', '?'):>4}] {para.get('env_label', ''):<22} -> {chunk_type}")
        {
            "abstract": lambda: self.render_abstract(doc, para),
            "section": lambda: self.render_heading(doc, para, 1),
            "subsection": lambda: self.render_heading(doc, para, 2),
            "subsubsection": lambda: self.render_heading(doc, para, 3),
            "equation": lambda: self.render_equation(doc, para),
            "figure": lambda: self.render_figure(doc, para, figures_dir),
            "table": lambda: self.render_table(doc, para),
            "code": lambda: self.render_code(doc, para),
            "mathenv": lambda: self.render_mathenv(doc, para),
            "proof": lambda: self.render_proof(doc, para),
            "inline_heading": lambda: self.render_inline_heading(doc, para),
            "paragraph": lambda: self.render_paragraph(doc, para),
        }.get(chunk_type, lambda: self.render_paragraph(doc, para))()
